import os
import re
import json
import uuid
import datetime
import streamlit as st
import asyncio
from fastmcp import Client
from PIL import Image
import base64
import streamlit.components.v1 as components
from streamlit_pdf_viewer import pdf_viewer
import pandas as pd



st.set_page_config(
    page_title="Assistente de Suporte Técnico",
    layout="wide",
    initial_sidebar_state="expanded"
)


# --- Lê o arquivo da fonte e converte em base64 ---
try:
    with open("fonts/Orbitron-Medium.ttf", "rb") as f:
        font_base64 = base64.b64encode(f.read()).decode()
except FileNotFoundError:
    font_base64 = ""  # fonte ausente não deve travar a aplicação

# --- Injeta a fonte no estilo global ---
st.markdown(f"""
    <style>
    @font-face {{
        font-family: 'Orbitron';
        src: url(data:font/ttf;base64,{font_base64}) format('truetype');
        font-weight: normal;
    }}
    h1 {{
        font-family: 'Orbitron', sans-serif !important;
    }}
    </style>
""", unsafe_allow_html=True)

# --- CSS principal ---
st.markdown("""
    <style>
        [data-testid="stSidebar"] { background-color: #4C2CCD !important; }
        footer[data-testid="stChatInput"] { background-color: #4C2CCD !important; }
        [data-testid="stChatInput"] textarea { background-color: #4C2CCD !important; color: white !important; }
        [data-testid="stChatInput"] label { color: white !important; }
        ::placeholder { color: white !important; }
        [data-testid="stAppViewContainer"] { background-color: white !important; color: black !important; }
        .stMarkdown, .stText, .stChatMessageContent { color: black !important; }
        .stMarkdown p, .stMarkdown span, .stMarkdown li { color: black !important; -webkit-text-fill-color: black !important; }
        .stChatMessage > div { background-color: #f5f5f5 !important; color: black !important; }
        .stChatMessage:nth-child(even) > div { background-color: #ffffff !important; }
        .stChatMessage p, .stChatMessage span { color: black !important; -webkit-text-fill-color: black !important; }
        /* File uploader nome/tamanho após seleção */
        section[data-testid="stSidebar"] .stSelectbox label {
            color: white !important;
            font-size: 1.2em !important;
            font-weight: 600 !important;
            letter-spacing: 0.5px;
        }
        section[data-testid="stSidebar"] div.stButton > button {
            color: white !important;
            background-color: #0a019e !important;
            border: 1px solid white !important;
        }
        .stFileUploader [data-testid="stFileUploaderFileName"] {
            color: white !important;
            font-weight: 500 !important;
        }
        .stFileUploader [data-testid="stFileUploaderFileSize"] {
            color: white !important;
        }
        .stFileUploader small {
            color: white !important;
        }
        /* Sidebar: todos os textos brancos (sessões, labels, captions) */
        section[data-testid="stSidebar"] p,
        section[data-testid="stSidebar"] span,
        section[data-testid="stSidebar"] label,
        section[data-testid="stSidebar"] small,
        section[data-testid="stSidebar"] h1,
        section[data-testid="stSidebar"] h2,
        section[data-testid="stSidebar"] h3,
        section[data-testid="stSidebar"] li,
        section[data-testid="stSidebar"] a {
            color: #ffffff !important;
            -webkit-text-fill-color: #ffffff !important;
        }
        section[data-testid="stSidebar"] button,
        section[data-testid="stSidebar"] button p,
        section[data-testid="stSidebar"] button span,
        section[data-testid="stSidebar"] button div {
            color: #ffffff !important;
            -webkit-text-fill-color: #ffffff !important;
        }

    </style>
""", unsafe_allow_html=True)

st.markdown(
    """
    <style>
    /* BASEWEB SELECT — texto branco globalmente */
    div[data-baseweb="select"] * {
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
    }
    div[data-baseweb="select"] input {
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
        caret-color: #FFFFFF !important;
    }
    div[data-baseweb="select"] input::placeholder {
        color: rgba(255,255,255,0.90) !important;
        -webkit-text-fill-color: rgba(255,255,255,0.90) !important;
    }
    div[data-baseweb="select"] svg {
        fill: #FFFFFF !important;
        color: #FFFFFF !important;
    }
    ul[role="listbox"] * {
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
    }
    input, textarea {
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
        caret-color: #FFFFFF !important;
    }
    input::placeholder, textarea::placeholder {
        color: rgba(255,255,255,0.75) !important;
        -webkit-text-fill-color: rgba(255,255,255,0.75) !important;
    }
    </style>
    """,
    unsafe_allow_html=True
)

st.markdown(
    """
    <style>
    /* FIX SIDEBAR: listbox com fundo branco — texto preto para aparecer */
    section[data-testid="stSidebar"] div[data-baseweb="select"] * {
        color: #000000 !important;
        -webkit-text-fill-color: #000000 !important;
        caret-color: #000000 !important;
    }
    section[data-testid="stSidebar"] div[data-baseweb="select"] input {
        color: #000000 !important;
        -webkit-text-fill-color: #000000 !important;
    }
    section[data-testid="stSidebar"] div[data-baseweb="select"] input::placeholder {
        color: rgba(0,0,0,0.60) !important;
        -webkit-text-fill-color: rgba(0,0,0,0.60) !important;
    }
    section[data-testid="stSidebar"] div[data-baseweb="select"] svg {
        fill: #000000 !important;
        color: #000000 !important;
    }
    section[data-testid="stSidebar"] ul[role="listbox"] * {
        color: #000000 !important;
        -webkit-text-fill-color: #000000 !important;
    }

    /* ── st.text_input na sidebar: fundo roxo escuro + texto branco ── */
    section[data-testid="stSidebar"] [data-baseweb="base-input"],
    section[data-testid="stSidebar"] [data-baseweb="base-input"] input {
        background-color: #2d1a9e !important;
        color: #ffffff !important;
        -webkit-text-fill-color: #ffffff !important;
        caret-color: #ffffff !important;
        border: 1px solid rgba(255,255,255,0.4) !important;
    }
    section[data-testid="stSidebar"] input[type="text"]:not([data-baseweb]) {
        background-color: #2d1a9e !important;
        color: #ffffff !important;
        -webkit-text-fill-color: #ffffff !important;
        caret-color: #ffffff !important;
    }
    section[data-testid="stSidebar"] input::placeholder {
        color: rgba(255,255,255,0.5) !important;
        -webkit-text-fill-color: rgba(255,255,255,0.5) !important;
    }
    </style>
    """,
    unsafe_allow_html=True
)

st.markdown(
    """
    <style>
    /* MULTISELECT CHIPS */
    [data-baseweb="tag"] {
        background-color: #0A019E !important;
        border: 1px solid rgba(255,255,255,0.55) !important;
    }
    [data-baseweb="tag"] * {
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
    }
    [data-baseweb="tag"] svg {
        fill: #FFFFFF !important;
        color: #FFFFFF !important;
    }
    div[data-baseweb="select"] [role="button"] {
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# CSS de alta especificidade para popover + JS como reforço
st.markdown("""
<style>
/* ── POPOVER: máxima especificidade com seletores encadeados ── */
div[data-testid="stPopoverContent"] {
    background-color: #3a1fb0 !important;
    border: 1px solid rgba(255,255,255,0.3) !important;
    border-radius: 8px !important;
}

/* Label "Renomear:" — seletor específico para label dentro do popover */
div[data-testid="stPopoverContent"] label,
div[data-testid="stPopoverContent"] label p,
div[data-testid="stPopoverContent"] label span,
div[data-testid="stPopoverContent"] [data-testid="stWidgetLabel"],
div[data-testid="stPopoverContent"] [data-testid="stWidgetLabel"] p {
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}

/* Todos os textos dentro do popover */
div[data-testid="stPopoverContent"] p,
div[data-testid="stPopoverContent"] span,
div[data-testid="stPopoverContent"] div,
div[data-testid="stPopoverContent"] small {
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}

/* Input: fundo roxo para combinar com o popover */
div[data-testid="stPopoverContent"] input[type="text"],
div[data-testid="stPopoverContent"] input,
div[data-testid="stPopoverContent"] textarea {
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
    caret-color: #ffffff !important;
    background-color: #3a1fb0 !important;
    border: 1px solid rgba(255,255,255,0.5) !important;
}

/* Contador "32/60" */
div[data-testid="stPopoverContent"] [data-testid="InputInstructions"],
div[data-testid="stPopoverContent"] [data-testid="InputInstructions"] * {
    color: rgba(255,255,255,0.7) !important;
    -webkit-text-fill-color: rgba(255,255,255,0.7) !important;
}

/* Botões */
div[data-testid="stPopoverContent"] button {
    background-color: #0a019e !important;
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
    border: 1px solid rgba(255,255,255,0.45) !important;
    border-radius: 6px !important;
}
div[data-testid="stPopoverContent"] button p,
div[data-testid="stPopoverContent"] button span,
div[data-testid="stPopoverContent"] button div {
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}
div[data-testid="stPopoverContent"] hr {
    border-color: rgba(255,255,255,0.25) !important;
}

/* Captions sidebar */
section[data-testid="stSidebar"] [data-testid="stCaptionContainer"] p,
section[data-testid="stSidebar"] small {
    color: rgba(255,255,255,0.75) !important;
    -webkit-text-fill-color: rgba(255,255,255,0.75) !important;
}
</style>
<script>
(function() {
    const ROXO = '#3a1fb0';
    const BTN  = '#0a019e';
    const W    = '#ffffff';

    // ── Injeta stylesheet no <head> para listbox ─────────────────────────
    // Técnica: criar um <style> element e inserir no head.
    // Seletores de atributo específicos vencem inline styles do BaseWeb
    // porque o navegador aplica inline styles ANTES de checar !important em stylesheets.
    // A exceção: quando usamos a pseudo-classe :is() ou ID único, aumentamos especificidade.
    // Mas a forma mais confiável é: stylesheet injetada + setProperty via JS juntos.
    const styleEl = document.createElement('style');
    styleEl.id = 'zilla-listbox-styles';
    styleEl.textContent = [
        'ul[role="listbox"] { background-color: #ffffff !important; }',
        'ul[role="listbox"] li { background-color: #ffffff !important; color: #000000 !important; }',
        'ul[role="listbox"] li * { color: #000000 !important; -webkit-text-fill-color: #000000 !important; background-color: transparent !important; }',
        'ul[role="listbox"] li:hover { background-color: ' + BTN + ' !important; }',
        'ul[role="listbox"] li:hover * { color: #ffffff !important; -webkit-text-fill-color: #ffffff !important; }'
    ].join(' ');
    document.head.appendChild(styleEl);

    function s(el, props) {
        Object.entries(props).forEach(([k,v]) => el.style.setProperty(k,v,'important'));
    }

    // Força inline também (dupla garantia)
    function paintListbox(ul) {
        s(ul, { 'background-color': '#ffffff' });
        ul.querySelectorAll('li, [role="option"]').forEach(li => {
            s(li, { 'background-color': '#ffffff', 'color': '#000000', '-webkit-text-fill-color': '#000000' });
            li.querySelectorAll('*').forEach(c =>
                s(c, { 'color': '#000000', '-webkit-text-fill-color': '#000000', 'background-color': 'transparent' })
            );
        });
    }

    // setInterval enquanto listbox estiver aberto
    let iv = null;
    function watchListboxes() {
        const uls = document.querySelectorAll('ul[role="listbox"]');
        if (!uls.length) { clearInterval(iv); iv = null; return; }
        uls.forEach(paintListbox);
    }

    // ── Popover de sessões ───────────────────────────────────────────────
    const popoverWatched = new WeakSet();
    const inputsHijacked = new WeakSet();

    // Intercepta o setter de style.color no elemento para impedir que
    // o React/Streamlit reescreva a cor depois do nosso setProperty
    function hijackInputColor(el) {
        if (inputsHijacked.has(el)) return;
        inputsHijacked.add(el);

        // Streamlit pode usar setAttribute('style', '...') ou el.style.xxx = '...'
        // Interceptamos os dois caminhos

        // 1) Intercepta setAttribute para filtrar background branco no style
        const origSetAttr = el.setAttribute.bind(el);
        el.setAttribute = function(name, value) {
            if (name === 'style' && typeof value === 'string') {
                value = value
                    .replace(/background-color[ ]*:[ ]*(white|#fff(?:fff)?|rgb[(]255,[ ]*255,[ ]*255[)])[^;]*;?/gi,
                             'background-color: #3a1fb0;')
                    .replace(/color[ ]*:[ ]*(black|#000(?:000)?|rgb[(]0,[ ]*0,[ ]*0[)])[^;]*;?/gi,
                             'color: #ffffff;');
            }
            return origSetAttr(name, value);
        };

        // 2) Intercepta el.style.backgroundColor e el.style.color via property descriptor
        try {
            const desc_bg = Object.getOwnPropertyDescriptor(CSSStyleDeclaration.prototype, 'backgroundColor');
            if (desc_bg && desc_bg.set) {
                Object.defineProperty(el.style, 'backgroundColor', {
                    get: desc_bg.get ? desc_bg.get.bind(el.style) : undefined,
                    set(v) {
                        const vs = String(v).replace(/ /g,'').toLowerCase();
                        if (['white','#fff','#ffffff','rgb(255,255,255)'].includes(vs)) return;
                        desc_bg.set.call(this, v);
                    },
                    configurable: true
                });
            }
        } catch(e) {}

        try {
            const desc_c = Object.getOwnPropertyDescriptor(CSSStyleDeclaration.prototype, 'color');
            if (desc_c && desc_c.set) {
                Object.defineProperty(el.style, 'color', {
                    get: desc_c.get ? desc_c.get.bind(el.style) : undefined,
                    set(v) {
                        const vs = String(v).replace(/ /g,'').toLowerCase();
                        if (['black','#000','#000000','rgb(0,0,0)'].includes(vs)) return;
                        desc_c.set.call(this, v);
                    },
                    configurable: true
                });
            }
        } catch(e) {}

        // Aplica estilos corretos
        el.style.setProperty('color', '#ffffff', 'important');
        el.style.setProperty('-webkit-text-fill-color', '#ffffff', 'important');
        el.style.setProperty('caret-color', '#ffffff', 'important');
        el.style.setProperty('background-color', '#3a1fb0', 'important');
        el.style.setProperty('border', '1px solid rgba(255,255,255,0.5)', 'important');
        el.style.setProperty('border-radius', '4px', 'important');
    }

    function paintPopover(p) {
        s(p, { 'background-color': ROXO, 'border': '1px solid rgba(255,255,255,0.3)', 'border-radius': '8px' });
        p.querySelectorAll('label, label *, p, span, div, small, li, a').forEach(el =>
            s(el, { 'color': W, '-webkit-text-fill-color': W }));
        p.querySelectorAll('input, textarea').forEach(el => hijackInputColor(el));
        p.querySelectorAll('button, button *').forEach(el =>
            s(el, { 'background-color': BTN, 'color': W, '-webkit-text-fill-color': W,
                'border': '1px solid rgba(255,255,255,0.45)', 'border-radius': '6px' }));
        p.querySelectorAll('hr').forEach(el => s(el, { 'border-color': 'rgba(255,255,255,0.25)' }));
    }

    function watchPopover(p) {
        if (popoverWatched.has(p)) return;
        popoverWatched.add(p);
        paintPopover(p);
        const iv = setInterval(() => {
            if (!document.body.contains(p)) { clearInterval(iv); return; }
            paintPopover(p);
        }, 50);
        let painting = false;
        new MutationObserver(() => {
            if (painting) return;
            painting = true;
            paintPopover(p);
            Promise.resolve().then(() => { painting = false; });
        }).observe(p, { childList: true, subtree: true, attributes: true, attributeFilter: ['style', 'class'] });
    }

    new MutationObserver(() => {
        const uls = document.querySelectorAll('ul[role="listbox"]');
        if (uls.length && !iv) {
            uls.forEach(paintListbox);
            iv = setInterval(watchListboxes, 16);
        }
        document.querySelectorAll('[data-testid="stPopoverContent"]').forEach(watchPopover);
    }).observe(document.body, { childList: true, subtree: true });

    document.querySelectorAll('[data-testid="stPopoverContent"]').forEach(watchPopover);

    // Renomeia e estiliza o botão "Browse files"
    function renameBrowseBtn() {
        document.querySelectorAll('.stFileUploader button').forEach(btn => {
            if (btn.textContent.trim().startsWith('Browse') || btn.textContent.trim() === 'Browse files') {
                btn.textContent = '📎 Enviar Arquivos';
            }
            // Força cor independente do texto
            btn.style.setProperty('background-color', '#0a019e', 'important');
            btn.style.setProperty('background',       '#0a019e', 'important');
            btn.style.setProperty('color',            '#ffffff', 'important');
            btn.style.setProperty('-webkit-text-fill-color', '#ffffff', 'important');
            btn.style.setProperty('border',           '1px solid #ffffff', 'important');
            btn.style.setProperty('border-radius',    '8px', 'important');
            btn.style.setProperty('width',            '100%', 'important');
            btn.style.setProperty('font-size',        '1rem', 'important');
            btn.style.setProperty('font-weight',      '500', 'important');
            btn.style.setProperty('padding',          '10px', 'important');
        });
    }
    setInterval(renameBrowseBtn, 500);
    renameBrowseBtn();
})();
</script>
""", unsafe_allow_html=True)



# -------------------------------------------------------------------
# MCP
# -------------------------------------------------------------------
# MCP — configuração e camada de resiliência
# -------------------------------------------------------------------
MCP_SERVER_URL = "http://localhost:8002/sse"

# Timeout por tipo de operação (segundos).
#
# index_document com imagens (OCR via qwen2.5vl:7b) pode levar vários minutos
# em hardware com recursos limitados — especialmente fluxogramas BPMN densos.
# ask_question demora mais por causa do reranking (N chamadas ao LLM).
# As demais ferramentas são operações rápidas (consultas ao Qdrant/filesystem).
MCP_TIMEOUT_FAST    = 30    # delete_document, feedback, list_*, confirmar_resposta
MCP_TIMEOUT_INDEX   = 600   # index_document: OCR de imagens complexas pode demorar ~5-10 min
MCP_TIMEOUT_SLOW    = 300   # ask_question: reranking + geração de resposta

# Número máximo de tentativas automáticas em caso de ConnectionClosed / McpError.
# Para index_document NÃO retentar automaticamente — o OCR é idempotente mas custoso;
# uma segunda tentativa imediata apenas piora a sobrecarga do hardware.
MCP_MAX_RETRIES       = 3
MCP_MAX_RETRIES_INDEX = 1   # sem retry para indexação (OCR pesado)
# Pausa base entre tentativas (segundos); dobra a cada tentativa (backoff exponencial)
MCP_RETRY_BACKOFF = 2.0

# Tools que exigem timeout longo (geração + reranking ou OCR)
_SLOW_TOOLS  = {"ask_question", "gerar_nome_sessao"}
_INDEX_TOOLS = {"index_document"}


def run_async(coro):
    """
    Executa uma coroutine de forma segura tanto dentro quanto fora de um event loop.

    Se já houver um loop ativo (cenário comum no Streamlit ≥ 1.18), executa em
    thread separada via ThreadPoolExecutor para não causar "event loop already running".
    """
    import concurrent.futures
    try:
        asyncio.get_running_loop()
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
            future = pool.submit(asyncio.run, coro)
            return future.result()
    except RuntimeError:
        return asyncio.run(coro)


async def _call_tool_once(tool_name: str, timeout: float, **kwargs):
    """
    Tentativa única de chamada a uma tool MCP com timeout forçado diretamente
    no httpcore — a camada onde o ReadTimeout realmente se origina.

    Análise do traceback:
      httpcore/_backends/anyio.py → read() → ReadTimeout
      ↑ essa é a origem real, não o httpx.AsyncClient.
      O httpcore recebe o timeout via objeto `Timeouts` passado em cada request.
      O mcp/client/sse.py cria internamente:
        httpx.AsyncClient() → sem timeout explícito
        httpx usa httpcore com timeout=DEFAULT (5s de read)

    Solução — patch em httpcore.AsyncConnectionPool.handle_async_request:
      Interceptamos a chamada no nível do httpcore e injetamos um Timeouts
      com read_timeout=timeout antes de cada request SSE. O patch é revertido
      no finally, garantindo que não vaza para outras chamadas.
    """
    import httpx
    import httpcore
    from fastmcp import Client as FastMCPClient

    # ----------------------------------------------------------------
    # Patch no httpcore: intercepta handle_async_request e força o timeout
    # ----------------------------------------------------------------
    _orig_handle = httpcore.AsyncConnectionPool.handle_async_request

    async def _patched_handle(self, request):
        # Substitui os timeouts do request pelo valor correto
        request.extensions["timeout"] = {
            "connect": 10.0,
            "read":    timeout,   # ← aqui é onde o ReadTimeout se originava
            "write":   30.0,
            "pool":    10.0,
        }
        return await _orig_handle(self, request)

    httpcore.AsyncConnectionPool.handle_async_request = _patched_handle
    try:
        async with FastMCPClient(MCP_SERVER_URL) as client:
            return await asyncio.wait_for(
                client.call_tool(tool_name, kwargs),
                timeout=timeout + 30,  # camada extra além do httpcore
            )
    finally:
        httpcore.AsyncConnectionPool.handle_async_request = _orig_handle


async def call_tool_async(tool_name: str, **kwargs):
    """
    Chama uma tool MCP com retry automático e backoff exponencial.

    Timeouts por categoria:
      - index_document  → MCP_TIMEOUT_INDEX (OCR pesado, sem retry automático)
      - ask_question    → MCP_TIMEOUT_SLOW  (reranking + LLM)
      - demais          → MCP_TIMEOUT_FAST

    Trata especificamente:
      - McpError: Connection closed  → servidor fechou a conexão SSE (timeout/crash)
      - asyncio.TimeoutError         → operação demorou mais que o limite configurado
      - ConnectionRefusedError       → servidor não está rodando
      - Exception genérica           → qualquer outro erro de transporte
    """
    import mcp.shared.exceptions as mcp_exc
    import httpx

    # Seleciona timeout e limite de retries conforme o tipo de tool
    if tool_name in _INDEX_TOOLS:
        timeout     = MCP_TIMEOUT_INDEX
        max_retries = MCP_MAX_RETRIES_INDEX   # OCR pesado: sem retry automático
    elif tool_name in _SLOW_TOOLS:
        timeout     = MCP_TIMEOUT_SLOW
        max_retries = MCP_MAX_RETRIES
    else:
        timeout     = MCP_TIMEOUT_FAST
        max_retries = MCP_MAX_RETRIES

    last_exc = None

    for attempt in range(1, max_retries + 1):
        try:
            return await _call_tool_once(tool_name, timeout, **kwargs)

        except mcp_exc.McpError as e:
            last_exc = e
            msg = str(e).lower()
            if "connection closed" in msg or "connection reset" in msg:
                reason = "conexão SSE fechada pelo servidor"
            else:
                reason = f"erro de protocolo MCP: {e}"

        except asyncio.TimeoutError:
            last_exc = asyncio.TimeoutError(f"Timeout de {timeout}s excedido para '{tool_name}'")
            reason   = f"timeout ({timeout}s) — servidor sobrecarregado ou OCR muito lento"

        except (httpx.ReadTimeout, httpx.WriteTimeout, httpx.PoolTimeout, httpx.TimeoutException) as e:
            # httpx.ReadTimeout é o erro que ocorria: servidor demorou mais que o read_timeout do httpx
            last_exc = e
            reason   = f"httpx timeout ({type(e).__name__}) — conexão HTTP expirou antes do servidor responder"

        except ConnectionRefusedError as e:
            raise ConnectionRefusedError(
                f"Servidor MCP recusou conexão em {MCP_SERVER_URL}. "
                "Verifique se o servidor está rodando com: python mcp_server.py"
            ) from e

        except Exception as e:
            last_exc = e
            reason   = f"erro inesperado: {type(e).__name__}: {e}"

        if attempt < max_retries:
            wait = MCP_RETRY_BACKOFF * (2 ** (attempt - 1))
            print(f"⚠️  Tentativa {attempt}/{max_retries} falhou para '{tool_name}' — {reason}. "
                  f"Aguardando {wait:.0f}s antes de retentar...")
            await asyncio.sleep(wait)
        else:
            print(f"❌ '{tool_name}' falhou na tentativa {attempt}/{max_retries} — {reason}. Sem mais retries.")

    raise RuntimeError(
        f"'{tool_name}' falhou após {max_retries} tentativa(s). "
        f"Último erro: {last_exc}"
    ) from last_exc


async def list_tools_async():
    """Lista as tools disponíveis no servidor MCP (sem retry — usada só na inicialização)."""
    async with Client(MCP_SERVER_URL) as client:
        return await asyncio.wait_for(client.list_tools(), timeout=MCP_TIMEOUT_FAST)

def extract_sources_from_response(resp_text: str) -> list[str]:
    """Extrai o bloco '📚 Fontes consultadas' da resposta do servidor."""
    if not resp_text:
        return []
    marker = "📚 **Fontes consultadas:**"
    if marker not in resp_text:
        return []
    try:
        _, sources_section = resp_text.split(marker, 1)
        lines = [l.strip() for l in sources_section.splitlines() if l.strip()]
        out = []
        for ln in lines:
            # Ex.: "1. 📄 arquivo.pdf (manual)"
            ln = re.sub(r"^\d+\.\s*📄\s*", "", ln).strip()
            if ln:
                out.append(ln)
        return out
    except Exception:
        return []

def extract_filenames(source_lines: list[str]) -> list[str]:
    """Converte 'arquivo.pdf (manual)' -> 'arquivo.pdf'."""
    out = []
    for s in source_lines or []:
        out.append(s.split(" (", 1)[0].strip())
    # remove duplicados preservando ordem
    seen = set()
    final = []
    for x in out:
        if x and x not in seen:
            seen.add(x)
            final.append(x)
    return final

def extract_types(source_lines: list[str]) -> list[str]:
    """Extrai tipos mostrados no texto 'arquivo.pdf (tipo)' -> ['manual', ...]."""
    types_found = []
    for s in source_lines or []:
        if "(" in s and s.endswith(")"):
            t = s.rsplit("(", 1)[-1].rstrip(")").strip().lower()
            if t:
                types_found.append(t)
    return sorted(set(types_found))


def _infer_feedback_from_text(free_text: str, rerank_scores: dict, mode: str):
    """
    Infere prefer_sources, avoid_sources, avoid_document_types, must_keywords
    e response_instruction a partir do texto livre do usuário + rerank_scores.

    Três categorias de feedback detectadas:
    1. EVITAR fonte/tipo: texto contém negação ("não use", "evite", etc.)
       → preenche avoid_sources / avoid_document_types
    2. INSTRUÇÃO DE FORMATO: texto descreve COMO responder ("em uma frase",
       "de forma objetiva", "use tópicos", "seja conciso", etc.)
       → preenche response_instruction, que é injetada no system_prompt do LLM
    3. KEYWORDS POSITIVAS: texto sem negação, sem instrução de formato
       → preenche must_keywords para expandir a query de busca
    """
    avoid_srcs            = []
    avoid_doc_types       = []
    prefer_srcs           = []
    must_kws              = []
    response_instruction  = ""

    text_lower = (free_text or "").lower()

    # Palavras/frases que indicam intenção de EVITAR algo
    negation_keywords = [
        "não use", "nao use", "não utilize", "nao utilize",
        "evite", "evitar", "sem usar", "excluir", "exclua",
        "ignore", "ignorar", "remover", "remova", "não considere",
        "nao considere", "não inclua", "nao inclua", "não usar",
        "nao usar",
    ]
    has_negation = any(kw in text_lower for kw in negation_keywords)

    if has_negation:
        # ------------------------------------------------------------------
        # Extrai avoid_sources: busca nome de arquivo mencionado no texto
        # ------------------------------------------------------------------
        for src in rerank_scores.keys():
            src_normalized = src.lower().replace("_", " ").replace("-", " ")
            src_no_ext     = re.sub(r"\.[a-z]{2,4}$", "", src_normalized).strip()

            if (src.lower()    in text_lower or
                src_normalized in text_lower or
                src_no_ext     in text_lower):
                avoid_srcs.append(src)
                print(f"[feedback] avoid_source inferido: '{src}'")

        # Fallback por score baixo se nenhuma fonte foi nomeada
        if not avoid_srcs:
            avoid_srcs = [s for s, sc in rerank_scores.items() if sc < 5.0]
            if avoid_srcs:
                print(f"[feedback] avoid_source por score baixo: {avoid_srcs}")

        # ------------------------------------------------------------------
        # [FIX-C] Extrai avoid_document_types: detecta tipo mencionado
        # junto de negação (ex.: "não use Manual", "evite relatórios antigos")
        # ------------------------------------------------------------------
        type_keywords = {
            "manual":      "manual",
            "manuais":     "manual",
            "relatório":   "relatório",
            "relatorio":   "relatório",
            "relatórios":  "relatório",
            "relatorios":  "relatório",
            "fluxograma":  "fluxograma",
            "fluxogramas": "fluxograma",
            "artigo":      "artigo",
            "artigos":     "artigo",
            "planilha":    "planilha",
            "planilhas":   "planilha",            
        }
        for word, doc_type in type_keywords.items():
            if word in text_lower and doc_type not in avoid_doc_types:
                avoid_doc_types.append(doc_type)
                print(f"[feedback] avoid_document_type inferido: '{doc_type}'")

    else:
        # Sem negação: extrai keywords técnicas positivas do texto livre do usuário.
        # [FIX-4] Regex corrigido: \w com re.UNICODE cobre letras acentuadas PT-BR
        # (ã, õ, á, é, í, ó, ú, â, ê, ô, etc.). O padrão anterior [a-zA-Z0-9_.\-]
        # rejeitava silenciosamente qualquer palavra com acento, fazendo must_keywords
        # ficar sempre vazia para feedback em português — a expansão de query no
        # fallback nunca ativava. Agora palavras como "memória", "módulo", "degradação"
        # são corretamente capturadas.
        stopwords = {"para", "este", "essa", "esse", "como", "qual", "quais",
                     "onde", "quando", "mais", "menos", "sobre", "entre",
                     "pela", "pelo", "numa", "neste", "nessa", "isso", "isto",
                     "falta", "faltou", "estava", "deveria", "resposta", "usou"}
        must_kws = [
            w.strip() for w in re.split(r"[\s,;:]+", free_text)
            if len(w.strip()) > 4
            and re.match(r"^[\w.\-]+$", w.strip(), re.UNICODE)  # [FIX-4] UNICODE
            and w.strip().lower() not in stopwords
        ]

    # Detecta menção POSITIVA de fonte no texto livre:
    # "use o Manual_do_RMA.pdf", "utilize o csv", "busque no manutencao memorias"
    # Funciona mesmo sem negação — complementa o ramo has_negation acima.
    positive_source_keywords = [
        "use ", "utilize ", "usar ", "utilizar ", "busque ", "buscar ",
        "prefira ", "preferir ", "consulte ", "consultar ", "do arquivo ",
        "do documento ", "a partir do ", "com base no ", "com base na ",
        "vem do ", "vem da ", "deve vir do ", "deve vir da ",
    ]
    has_positive_source = any(kw in text_lower for kw in positive_source_keywords)

    explicit_prefer_srcs = []
    if has_positive_source and not has_negation:
        for src in rerank_scores.keys():
            src_normalized = src.lower().replace("_", " ").replace("-", " ")
            src_no_ext     = re.sub(r"[.][a-z]{2,4}$", "", src_normalized).strip()
            # Tokens individuais: "manutencao memorias" bate em "manutencao_memorias_500.csv"
            src_tokens     = set(re.split(r"[\s_.\-]+", src_no_ext)) - {""}
            text_tokens    = set(re.split(r"[\s,;:.]+", text_lower)) - {""}
            token_overlap  = src_tokens & text_tokens
            matched = (
                src.lower()    in text_lower or
                src_normalized in text_lower or
                src_no_ext     in text_lower or
                # Match parcial: pelo menos 2 tokens do nome do arquivo aparecem no texto
                len(token_overlap) >= min(2, len(src_tokens))
            )
            if matched:
                explicit_prefer_srcs.append(src)
                print(f"[feedback] prefer_source explícito inferido: '{src}' (tokens={token_overlap})")

    # Combina: fontes explicitamente mencionadas + fontes com score alto
    prefer_srcs = list({
        *explicit_prefer_srcs,
        *[s for s, sc in rerank_scores.items() if sc >= 6.0 and s not in avoid_srcs]
    })

    # [FIX-INSTRUCAO] Detecta instrução de formato no texto livre.
    # Ativado quando NÃO há negação (não é sobre evitar fontes) e o texto
    # contém palavras que descrevem COMO responder (formato, tom, tamanho).
    # O texto inteiro é armazenado como response_instruction para ser
    # injetado no system_prompt do LLM em perguntas similares futuras.
    format_keywords = [
        "frase", "frases", "objetivo", "objetiva", "objetivamente",
        "conciso", "concisa", "resumo", "resumida", "resumido",
        "simples", "direta", "direto", "curto", "curta",
        "tópicos", "topicos", "lista", "detalhado", "detalhada",
        "explicar melhor", "mais detalhes", "parágrafo", "paragrafo",
        "formal", "informal", "técnico", "técnica",
    ]
    if (
        not has_negation
        and free_text
        and any(kw in (free_text or "").lower() for kw in format_keywords)
    ):
        response_instruction = free_text.strip()

    return prefer_srcs, avoid_srcs, avoid_doc_types, must_kws, response_instruction, explicit_prefer_srcs

# -------------------------------------------------------------------
# Conecta no servidor MCP (mantendo seu padrão)
# -------------------------------------------------------------------
try:
    server_tools = run_async(list_tools_async())
except Exception as e:
    st.error(f"⚠️ Nesse momento não foi possível conectar ao servidor MCP em {MCP_SERVER_URL}: {e}")
    st.info("Verifique se o servidor MCP está rodando com: `python mcp_server.py`")
    st.stop()

# -------------------------------------------------------------------
# Header
# -------------------------------------------------------------------
st.markdown(
    "<h1 style='text-align: center; font-size: 2.5em; font-family: Orbitron, sans-serif; color: #4C2CCD;'>🤖 Assistente de Suporte Técnico</h1>",
    unsafe_allow_html=True
)

st.markdown("<h3 style='text-align: center; font-size: 1em; color: #4C2CCD;'>Essa IA responde dúvidas com base em documentos enviados sobre o processo RMA, funcionamento do sistema iRMA e outros.</h3>", unsafe_allow_html=True)
st.markdown("<h3 style='text-align: center; font-size: 1em; color: #000000;'>Exemplos: O que é RMA? Quais RMAs envolvem ECC? <br> Liste os casos relacionados a temperatura e descreva evidências que suportam isso.<br> Quais são as etapas do processo usuário acessar sistema RMA? </h3>", unsafe_allow_html=True)

# -------------------------------------------------------------------
# Storage local de uploads (mantido)
# -------------------------------------------------------------------
UPLOAD_DIR   = "uploads_mcp"

# =============================================================================
# [SESSÕES] Funções auxiliares de gerenciamento de sessões
# =============================================================================

def _new_session_id() -> str:
    """Gera um UUID v4 como string, compatível com Qdrant (aceita UUID como point id)."""
    return str(uuid.uuid4())


def _session_state_init():
    """Garante que todas as chaves de sessão estejam no session_state."""
    defaults = {
        "messages":               [{"role": "assistant", "content": "Olá! Como posso ajudar você hoje?"}],
        "ultima_pergunta":        "",
        "ultima_resposta":        "",
        "ultima_fontes":          [],
        "show_confirm_button":    False,
        "force_fresh":            False,
        "show_feedback_expander": False,
        "feedback_saved":         False,
        "rerun_after_feedback":   False,
        "rerank_scores":          {},
        "feedback_mode":          "partial",
        "confirm_clear_knowledge": False,  # controla o passo de confirmação do botão limpar
        # Sessão ativa
        "session_id":             None,   # UUID da sessão persistida (None = sessão nova não salva)
        "session_name":           "",     # nome exibido na sidebar
        "session_created_at":     "",     # ISO 8601
        "session_name_editing":   False,  # controla o campo inline de renomear
        "sessions_list":          None,   # cache da lista de sessões (recarregado ao mudar)
    }
    for key, val in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = val


def _save_current_session(force_name: str = ""):
    """
    Persiste a sessão atual no Qdrant via tool MCP salvar_sessao.
    Se session_id for None (sessão nova), gera um novo UUID.
    Se force_name for informado, usa esse nome em vez do session_name atual.
    Atualiza session_state após o salvamento.
    """
    msgs = [
        m for m in st.session_state.messages
        if m["role"] in ("user", "assistant")
    ]
    if not msgs or all(m["role"] == "assistant" for m in msgs):
        return  # nada para salvar

    if not st.session_state.session_id:
        st.session_state.session_id     = _new_session_id()
        st.session_state.session_created_at = datetime.datetime.now(datetime.timezone.utc).isoformat()

    name = force_name or st.session_state.session_name or "Nova sessão"

    try:
        result = run_async(call_tool_async(
            "salvar_sessao",
            session_id=st.session_state.session_id,
            name=name,
            messages=json.dumps(msgs, ensure_ascii=False),
            pinned=False,
            created_at=st.session_state.session_created_at,
        ))
        msg = getattr(result, "data", getattr(result, "output", str(result)))
        if str(msg).strip() == "ok":
            st.session_state.session_name    = name
            st.session_state.sessions_list   = None  # invalida cache
    except Exception as e:
        print(f"⚠️ Falha ao salvar sessão: {e}")


def _load_sessions() -> list:
    """
    Carrega lista de sessões do servidor MCP.
    Usa cache em session_state.sessions_list para evitar chamadas repetidas.
    """
    if st.session_state.sessions_list is not None:
        return st.session_state.sessions_list
    try:
        result = run_async(call_tool_async("listar_sessoes"))
        raw    = getattr(result, "data", getattr(result, "output", str(result)))
        data   = json.loads(str(raw))
        if isinstance(data, list):
            st.session_state.sessions_list = data
            return data
    except Exception as e:
        print(f"⚠️ Falha ao listar sessões: {e}")
    return []


def _export_txt(session_name: str, messages: list) -> bytes:
    """Gera exportação da sessão em formato TXT."""
    lines = [
        f"SESSÃO: {session_name}",
        f"Exportado em: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}",
        "=" * 60,
        "",
    ]
    for msg in messages:
        role   = "👤 Usuário" if msg["role"] == "user" else "🤖 Assistente"
        lines.append(f"{role}:")
        lines.append(msg["content"])
        lines.append("")
    return "\n".join(lines).encode("utf-8")


def _export_docx(session_name: str, messages: list) -> bytes:
    """Gera exportação da sessão em DOCX via python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io

        doc = Document()

        for section in doc.sections:
            section.top_margin    = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

        # Título
        title = doc.add_heading(level=1)
        run = title.add_run(session_name)
        run.font.color.rgb = RGBColor(0x2E, 0x3A, 0x8C)
        run.font.size = Pt(16)

        # Data
        meta = doc.add_paragraph()
        mr = meta.add_run(f"Exportado em {datetime.datetime.now().strftime('%d/%m/%Y às %H:%M')}")
        mr.font.size = Pt(9)
        mr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_paragraph("─" * 60)

        for msg in messages:
            # Role
            role_par = doc.add_paragraph()
            role_run = role_par.add_run(
                "Usuário" if msg["role"] == "user" else "Assistente"
            )
            role_run.bold = True
            role_run.font.size = Pt(10)
            role_run.font.color.rgb = (
                RGBColor(0x0A, 0x01, 0x9E) if msg["role"] == "user"
                else RGBColor(0x1A, 0x7A, 0x4A)
            )

            # Conteúdo — cada linha como parágrafo
            for line in msg["content"].split("\n"):
                p = doc.add_paragraph()
                r = p.add_run(line)
                r.font.size = Pt(10)

            doc.add_paragraph("· " * 30)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except ImportError:
        return _export_txt(session_name, messages)
os.makedirs(UPLOAD_DIR, exist_ok=True)

try:
    image = Image.open('zilia_logo.png')
    st.sidebar.image(image, width='stretch')
except FileNotFoundError:
    st.sidebar.markdown("<span style='color:white; font-weight:bold;'>🤖 Assistente Técnico</span>", unsafe_allow_html=True)

doc_type_options = ["Relatório","Manual", "Fluxograma", "Artigo", "Planilha", "Outro"]
selected_doc_type = st.sidebar.selectbox("Para envio de documentos escolha um dos tipos abaixo:", doc_type_options, index=0)

# CSS agressivo aplicado via st.markdown IMEDIATAMENTE antes do file_uploader
# para garantir que o seletor tem alta especificidade no momento do render
st.sidebar.markdown(
    "<p style='color:#000000 !important; -webkit-text-fill-color:#000000 !important; "
    "font-weight:600; margin-bottom:2px;'>📎 Enviar Arquivos</p>"
    "<p style='color:#ffffff !important; -webkit-text-fill-color:#ffffff !important; "
    "font-size:0.78rem; margin-top:0; margin-bottom:4px; opacity:0.85;'>"
    "Limite 200MB por arquivo • PDF, PNG, JPG, JPEG, CSV</p>",
    unsafe_allow_html=True
)
st.sidebar.markdown("""
<style>
div[data-testid="stFileUploadDropzone"] {
    background: transparent !important;
    border: none !important;
    padding: 0 !important;
    margin: 0 !important;
}
div[data-testid="stFileUploaderDropzoneInstructions"] {
    display: none !important;
}
div[data-testid="stFileUploadDropzone"] button,
div[data-testid="stFileUploadDropzone"] button:hover,
div[data-testid="stFileUploadDropzone"] button:focus {
    background-color: #0a019e !important;
    background: #0a019e !important;
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
    border: 1px solid #ffffff !important;
    border-radius: 8px !important;
    width: 100% !important;
    padding: 10px !important;
    font-size: 1rem !important;
    font-weight: 500 !important;
    cursor: pointer !important;
}
</style>
""", unsafe_allow_html=True)
uploaded_file = st.sidebar.file_uploader(
    "upload",
    accept_multiple_files=False,
    type=["pdf", "png", "jpg", "jpeg", "docx", "txt", "csv"],
    label_visibility="collapsed",
)

if uploaded_file:
    if st.sidebar.button("📥 Enviar e Indexar Documento"):
        image_exts = {".png", ".jpg", ".jpeg"}
        text_exts  = {".docx", ".txt", ".csv"}
        ext      = os.path.splitext(uploaded_file.name)[1].lower()
        is_image = ext in image_exts
        is_text  = ext in text_exts

        if is_image:
            st.sidebar.warning(
                "⏳ **Imagem detectada.** O OCR via `qwen2.5vl:7b` pode levar "
                "**5 a 10 minutos** para imagens complexas (ex.: fluxogramas BPMN). "
                "Aguarde — não feche nem recarregue a página."
            )
        elif is_text:
            st.sidebar.info(
                f"📄 Arquivo de texto detectado ({ext.upper()}). "
                "A indexação será rápida — sem OCR necessário."
            )

        file_path = os.path.join(UPLOAD_DIR, uploaded_file.name)

        # Salva o arquivo localmente antes de enviar ao servidor
        try:
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
        except Exception as e:
            st.sidebar.error(f"❌ Erro ao salvar '{uploaded_file.name}' localmente: {e}")
            uploaded_file = None  # impede o bloco abaixo de executar

        if uploaded_file:
            spinner_msg = (
                f"🔍 Processando OCR de '{uploaded_file.name}' (pode demorar vários minutos)..."
                if is_image else
                f"📄 Indexando '{uploaded_file.name}'..."
            )

            indexing_succeeded = False

            with st.spinner(spinner_msg):
                try:
                    result = run_async(call_tool_async(
                        'index_document',
                        file_path=file_path,
                        filename=uploaded_file.name,
                        document_type=selected_doc_type
                    ))
                    msg = getattr(result, 'data', getattr(result, 'output', str(result)))
                    if not isinstance(msg, str):
                        msg = str(msg)

                    is_server_error = msg.startswith(("⚠️", "❌", "Erro", "Não foi"))
                    if is_server_error:
                        st.sidebar.error(f"⚠️ Servidor: {msg}")
                    else:
                        st.sidebar.success(f"✅ {msg}")
                        indexing_succeeded = True
                        st.session_state.knowledge_stats = None  # força atualização do painel

                except ConnectionRefusedError:
                    st.sidebar.error(
                        "⚠️ Servidor MCP fora do ar. "
                        "Inicie com: `python mcp_server.py`"
                    )

                except RuntimeError as e:
                    err_str = str(e).lower()
                    is_timeout_err = any(x in err_str for x in (
                        "timeout", "connection closed", "readtimeout", "httpx"
                    ))
                    if is_timeout_err:
                        if is_image:
                            st.sidebar.error(
                                f"⏱️ **Timeout no OCR de '{uploaded_file.name}'.**\n\n"
                                "O modelo `qwen2.5vl:7b` demorou mais do que o limite configurado. "
                                "Possíveis causas:\n"
                                "- Imagem muito complexa ou de alta resolução\n"
                                "- Hardware sem GPU ou memória RAM insuficiente\n\n"
                                "**Sugestões:**\n"
                                "- Reduza a resolução da imagem antes de enviar\n"
                                "- Converta o fluxograma para PDF (texto embutido é muito mais rápido)\n"
                                "- Aguarde o servidor liberar memória e tente novamente"
                            )
                        else:
                            st.sidebar.error(
                                f"⏱️ Timeout ao indexar '{uploaded_file.name}'. "
                                "O servidor pode estar sobrecarregado. Tente novamente."
                            )
                    else:
                        st.sidebar.error(f"❌ Erro ao indexar '{uploaded_file.name}': {e}")

                except Exception as e:
                    st.sidebar.error(
                        f"❌ Erro inesperado ao indexar '{uploaded_file.name}': "
                        f"{type(e).__name__}: {e}"
                    )

                finally:
                    if not indexing_succeeded and os.path.exists(file_path):
                        try:
                            os.remove(file_path)
                            st.sidebar.info(
                                f"🗑️ '{uploaded_file.name}' removido de uploads_mcp "
                                "(indexação não concluída)."
                            )
                            print(f"🗑️ Arquivo órfão removido pelo cliente: {file_path}")
                        except Exception as cleanup_err:
                            print(f"⚠️ Falha ao remover arquivo órfão '{file_path}': {cleanup_err}")

# ==========================
# Listagem e visualização
# ==========================
st.sidebar.markdown("<span style='color: white;'>**📂 Arquivos enviados**</span>", unsafe_allow_html=True)

try:
    available_files = [
        f for f in sorted(os.listdir(UPLOAD_DIR))
        if os.path.isfile(os.path.join(UPLOAD_DIR, f))
    ]
except Exception:
    available_files = []

if not available_files:
    st.sidebar.info("Nenhum arquivo encontrado em uploads_mcp")
    selected_view_file = None
else:
    selected_view_file = st.sidebar.selectbox(
        "Selecione um arquivo para visualizar:",
        options=["—"] + available_files,
        index=0
    )

if selected_view_file and selected_view_file != "—":
    view_path = os.path.join(UPLOAD_DIR, selected_view_file)
    ext = os.path.splitext(selected_view_file)[1].lower()

    st.markdown(
        f"<div style='background-color: white; padding: 10px; border-radius: 5px; margin: 10px 0;'><h3><span style='color: black;'>Visualizando:</span> <span style='color: #0a019e;'>{selected_view_file}</span></h3></div>",
        unsafe_allow_html=True
    )

    if ext == ".pdf":
        try:
            with open(view_path, "rb") as f:
                binary_data = f.read()
            pdf_viewer(input=binary_data, width=700)
        except Exception as e:
            st.error(f"Erro ao abrir PDF: {e}")
    elif ext in [".png", ".jpg", ".jpeg"]:
        try:
            img = Image.open(view_path)
            st.image(img, caption=selected_view_file, width='stretch')
        except Exception as e:
            st.error(f"Erro ao abrir imagem: {e}")
    elif ext in [".txt", ".csv"]:
        try:
            if ext == ".csv":
                for enc in ("utf-8", "latin-1"):
                    try:
                        df = pd.read_csv(view_path, encoding=enc)
                        break
                    except UnicodeDecodeError:
                        continue
                st.dataframe(df, use_container_width=True)
            else:
                for enc in ("utf-8", "latin-1"):
                    try:
                        with open(view_path, "r", encoding=enc) as f:
                            content = f.read()
                        break
                    except UnicodeDecodeError:
                        continue
                st.text_area("Conteúdo do arquivo:", value=content, height=400, disabled=True)
        except Exception as e:
            st.error(f"Erro ao abrir arquivo: {e}")
    elif ext == ".docx":
        st.info("📄 Arquivo Word (.docx) — visualização não disponível diretamente. Use o botão abaixo para baixar.")
        try:
            with open(view_path, "rb") as f:
                st.download_button(label="⬇️ Baixar arquivo DOCX", data=f, file_name=selected_view_file)
        except Exception as e:
            st.error(f"Erro ao preparar download: {e}")
    else:
        st.warning("Formato não suportado para visualização. Faça o download abaixo.")
        try:
            with open(view_path, "rb") as f:
                st.download_button(label="⬇️ Baixar arquivo", data=f, file_name=selected_view_file)
        except Exception as e:
            st.error(f"Erro ao preparar download: {e}")

    # ==========================
    # Botão de exclusão de documento
    # ==========================
    if st.sidebar.button("🗑️ Excluir documento selecionado", key=f"del_{selected_view_file}"):
        with st.spinner(f"Excluindo '{selected_view_file}'..."):
            try:
                delete_result = run_async(call_tool_async('delete_document', filename=selected_view_file))
                msg = getattr(delete_result, "data", getattr(delete_result, "output", str(delete_result)))
                st.success(f"{msg}")
                st.session_state.knowledge_stats = None  # força atualização do painel
                st.rerun()
            except Exception as e:
                st.error(f"Erro ao excluir documento: {e}")

# ==========================
# Inicialização do session_state
# ==========================
_session_state_init()

# ==========================
# Sidebar — Painel de Conhecimento
# ==========================
st.sidebar.markdown("---")
st.sidebar.markdown(
    "<span style='color:white; font-weight:bold; font-size:1.1em;'>🧠 Conhecimento acumulado</span>",
    unsafe_allow_html=True
)

# Carrega stats com cache em session_state para não chamar o servidor a cada rerun
if "knowledge_stats" not in st.session_state:
    st.session_state.knowledge_stats = None
if "knowledge_stats_error" not in st.session_state:
    st.session_state.knowledge_stats_error = False

# Botão de atualizar + carga automática na primeira vez
_stats_loaded = st.session_state.knowledge_stats is not None
_col_stats, _col_refresh = st.sidebar.columns([7, 3])

with _col_refresh:
    if st.button("🔄", key="btn_refresh_stats", help="Atualizar contadores", use_container_width=True):
        st.session_state.knowledge_stats = None  # força recarga
        st.rerun()

if not _stats_loaded:
    try:
        _stats_result = run_async(call_tool_async("get_knowledge_stats"))
        _stats_raw = getattr(_stats_result, "data", getattr(_stats_result, "output", str(_stats_result)))
        import json as _stats_json
        st.session_state.knowledge_stats = _stats_json.loads(str(_stats_raw))
        st.session_state.knowledge_stats_error = False
    except Exception as _stats_err:
        st.session_state.knowledge_stats_error = True
        print(f"⚠️ Erro ao carregar knowledge_stats: {_stats_err}")

_ks = st.session_state.knowledge_stats or {}
_err = st.session_state.knowledge_stats_error

if _err:
    st.sidebar.caption("⚠️ Não foi possível carregar os contadores.")
else:
    _cache_n    = _ks.get("cache_count", 0)
    _fb_n       = _ks.get("feedback_count", 0)
    _docs_n     = _ks.get("docs_count", 0)
    _chunks_n   = _ks.get("chunks_count", 0)

    # Linha 1: documentos e chunks
    _c1, _c2 = st.sidebar.columns(2)
    with _c1:
        st.sidebar.markdown(
            f"<div style='background:rgba(255,255,255,0.10); border-radius:8px; padding:8px 10px; margin-bottom:6px;'>"
            f"<div style='color:rgba(255,255,255,0.65); font-size:0.70em; text-transform:uppercase; letter-spacing:0.5px;'>📄 Documentos</div>"
            f"<div style='color:#ffffff; font-size:1.4em; font-weight:700; line-height:1.2;'>{_docs_n}</div>"
            f"</div>",
            unsafe_allow_html=True
        )
    with _c2:
        st.sidebar.markdown(
            f"<div style='background:rgba(255,255,255,0.10); border-radius:8px; padding:8px 10px; margin-bottom:6px;'>"
            f"<div style='color:rgba(255,255,255,0.65); font-size:0.70em; text-transform:uppercase; letter-spacing:0.5px;'>🧩 Trechos indexados</div>"
            f"<div style='color:#ffffff; font-size:1.4em; font-weight:700; line-height:1.2;'>{_chunks_n:,}</div>"
            f"</div>",
            unsafe_allow_html=True
        )

    # Linha 2: respostas confirmadas e preferências aprendidas
    _c3, _c4 = st.sidebar.columns(2)
    with _c3:
        st.sidebar.markdown(
            f"<div style='background:rgba(255,255,255,0.10); border-radius:8px; padding:8px 10px; margin-bottom:4px;'>"
            f"<div style='color:rgba(255,255,255,0.65); font-size:0.70em; text-transform:uppercase; letter-spacing:0.5px;'>✅ Respostas confirmadas</div>"
            f"<div style='color:#ffffff; font-size:1.4em; font-weight:700; line-height:1.2;'>{_cache_n}</div>"
            f"</div>",
            unsafe_allow_html=True
        )
    with _c4:
        st.sidebar.markdown(
            f"<div style='background:rgba(255,255,255,0.10); border-radius:8px; padding:8px 10px; margin-bottom:4px;'>"
            f"<div style='color:rgba(255,255,255,0.65); font-size:0.70em; text-transform:uppercase; letter-spacing:0.5px;'>🎯 Preferências aprendidas</div>"
            f"<div style='color:#ffffff; font-size:1.4em; font-weight:700; line-height:1.2;'>{_fb_n}</div>"
            f"</div>",
            unsafe_allow_html=True
        )

    # ── Botão de limpeza com confirmação em dois passos ──────────────────────
    if not st.session_state.get("confirm_clear_knowledge", False):
        # Passo 1: botão inicial — apenas mostra o alerta de confirmação
        if st.sidebar.button(
            "🗑️ Limpar conhecimento acumulado",
            key="btn_clear_knowledge",
            use_container_width=True,
            help="Remove respostas confirmadas e preferências aprendidas. Documentos e sessões não são afetados.",
        ):
            st.session_state.confirm_clear_knowledge = True
            st.rerun()
    else:
        # Passo 2: confirmação explícita antes de executar
        st.sidebar.markdown(
            "<div style='background:rgba(220,50,50,0.25); border:1px solid rgba(255,100,100,0.5); "
            "border-radius:8px; padding:10px; margin-bottom:8px;'>"
            "<span style='color:#ffcccc; font-size:0.85em;'>⚠️ Isso apagará todas as respostas "
            "confirmadas e preferências aprendidas. Documentos e sessões <b>não</b> serão afetados.</span>"
            "</div>",
            unsafe_allow_html=True,
        )
        _cc1, _cc2 = st.sidebar.columns(2)
        with _cc1:
            if st.button("✅ Confirmar", key="btn_clear_confirm", use_container_width=True):
                _clear_ok = False
                _clear_msgs = []
                _clear_err = None

                with st.sidebar:
                    with st.spinner("Limpando..."):
                        try:
                            _clear_result = run_async(call_tool_async(
                                "clear_learned_knowledge", target="all"
                            ))
                            _clear_raw = getattr(_clear_result, "data",
                                                 getattr(_clear_result, "output", str(_clear_result)))
                            import json as _cj
                            _cr = _cj.loads(str(_clear_raw))
                            _partial_fail = False
                            if _cr.get("cache", {}).get("ok"):
                                n = _cr["cache"].get("removidos", 0)
                                r = _cr["cache"].get("restantes", 0)
                                _clear_msgs.append(f"{n} resposta(s) confirmada(s)")
                                if r > 0:
                                    _partial_fail = True
                                    _clear_msgs.append(f"⚠️ {r} entrada(s) do cache não foram removidas")
                            if _cr.get("feedback", {}).get("ok"):
                                n = _cr["feedback"].get("removidos", 0)
                                r = _cr["feedback"].get("restantes", 0)
                                _clear_msgs.append(f"{n} preferência(s) aprendida(s)")
                                if r > 0:
                                    _partial_fail = True
                                    _clear_msgs.append(f"⚠️ {r} entrada(s) do feedback não foram removidas")
                            _clear_ok = not _partial_fail
                        except Exception as _ce:
                            _clear_err = str(_ce)

                # Atualiza session_state FORA do spinner/sidebar, antes do rerun
                # para garantir persistência antes do próximo ciclo de render
                st.session_state.confirm_clear_knowledge = False
                st.session_state.knowledge_stats = None  # força recarga dos contadores
                st.session_state.knowledge_stats_error = False

                if _clear_ok:
                    msg = f"🗑️ Removido: {' e '.join(_clear_msgs)}." if _clear_msgs else "Base já estava vazia."
                    st.session_state["_clear_success_msg"] = msg
                else:
                    st.session_state["_clear_success_msg"] = None
                    st.session_state["_clear_error_msg"] = _clear_err

                st.rerun()

        with _cc2:
            if st.button("❌ Cancelar", key="btn_clear_cancel", use_container_width=True):
                st.session_state.confirm_clear_knowledge = False
                st.rerun()

# Exibe mensagem de resultado da limpeza (após o rerun, fora do bloco de confirmação)
if st.session_state.get("_clear_success_msg"):
    st.sidebar.success(st.session_state.pop("_clear_success_msg"))
elif st.session_state.get("_clear_error_msg"):
    st.sidebar.error(f"❌ Erro ao limpar: {st.session_state.pop('_clear_error_msg')}")

# ==========================
# Sidebar — Sessões
# ==========================
st.sidebar.markdown("---")
st.sidebar.markdown("<span style='color:white; font-weight:bold; font-size:1.1em;'>💬 Sessões</span>", unsafe_allow_html=True)

# Botão Nova Sessão
if st.sidebar.button("➕ Nova Sessão", use_container_width=True):
    # Salva a sessão atual antes de limpar (se tiver conteúdo)
    if any(m["role"] == "user" for m in st.session_state.messages):
        _save_current_session()
    # Reinicia estado para nova sessão
    st.session_state.messages            = [{"role": "assistant", "content": "Olá! Como posso ajudar você hoje?"}]
    st.session_state.session_id          = None
    st.session_state.session_name        = ""
    st.session_state.session_created_at  = ""
    st.session_state.ultima_pergunta     = ""
    st.session_state.ultima_resposta     = ""
    st.session_state.ultima_fontes       = []
    st.session_state.show_confirm_button = False
    st.session_state.force_fresh         = False
    st.session_state.show_feedback_expander = False
    st.session_state.feedback_saved      = False
    st.session_state.rerank_scores       = {}
    st.session_state.sessions_list       = None
    st.rerun()

# Lista de sessões salvas
sessions = _load_sessions()

if not sessions:
    st.sidebar.caption("Nenhuma sessão salva ainda.")
else:
    # Separador visual entre pinadas e normais
    has_pinned   = any(s["pinned"] for s in sessions)
    has_unpinned = any(not s["pinned"] for s in sessions)

    if has_pinned:
        st.sidebar.markdown("<span style='color:rgba(255,255,255,0.6); font-size:0.8em;'>📌 FIXADAS</span>", unsafe_allow_html=True)

    for s in sessions:
        sid      = s["session_id"]
        sname    = s["name"]
        pinned   = s["pinned"]
        n_msgs   = s["message_count"]
        is_active = (sid == st.session_state.session_id)

        if not pinned and has_pinned and has_unpinned:
            # Separador entre pinadas e não-pinadas (só na primeira não-pinada)
            if s == next((x for x in sessions if not x["pinned"]), None):
                st.sidebar.markdown("<span style='color:rgba(255,255,255,0.6); font-size:0.8em;'>🕐 RECENTES</span>", unsafe_allow_html=True)

        # Nome com destaque se for a sessão ativa
        prefix  = "📌 " if pinned else ""
        label   = f"{prefix}{sname}"
        if is_active:
            label = f"▶ {label}"

        col_btn, col_act = st.sidebar.columns([7, 3])
        with col_btn:
            if st.button(label, key=f"sess_{sid}", use_container_width=True, help=f"{n_msgs} mensagens"):
                if sid != st.session_state.session_id:
                    # Salva sessão atual antes de trocar
                    if any(m["role"] == "user" for m in st.session_state.messages):
                        _save_current_session()
                    # Carrega a sessão selecionada
                    try:
                        result = run_async(call_tool_async("carregar_sessao", session_id=sid))
                        raw    = getattr(result, "data", getattr(result, "output", str(result)))
                        data   = json.loads(str(raw))
                        if "error" not in data:
                            st.session_state.messages           = data.get("messages", [])
                            st.session_state.session_id         = sid
                            st.session_state.session_name       = data.get("name", sname)
                            st.session_state.session_created_at = data.get("created_at", "")
                            st.session_state.ultima_pergunta    = ""
                            st.session_state.ultima_resposta    = ""
                            st.session_state.show_confirm_button = False
                            st.session_state.sessions_list      = None
                            st.rerun()
                        else:
                            st.sidebar.error(f"Erro: {data['error']}")
                    except Exception as e:
                        st.sidebar.error(f"Erro ao carregar sessão: {e}")

        with col_act:
            # Botão toggle que abre/fecha o menu inline na sidebar
            menu_key = f"menu_open_{sid}"
            if menu_key not in st.session_state:
                st.session_state[menu_key] = False

            toggle_label = "▲" if st.session_state[menu_key] else "⋮"
            if st.button(toggle_label, key=f"toggle_{sid}", use_container_width=True):
                st.session_state[menu_key] = not st.session_state[menu_key]
                st.rerun()

        # Menu inline — renderizado diretamente na sidebar (sem portal flutuante)
        if st.session_state.get(f"menu_open_{sid}", False):
            with st.sidebar:
                st.markdown(
                    "<div style='background:#3a1fb0; border:1px solid rgba(255,255,255,0.3); "
                    "border-radius:8px; padding:12px; margin:4px 0 8px 0;'>",
                    unsafe_allow_html=True,
                )
                new_name = st.text_input(
                    "Renomear:",
                    value=sname,
                    key=f"rename_{sid}",
                    max_chars=60,
                    label_visibility="visible",
                )
                if st.button("✏️ Salvar nome", key=f"save_name_{sid}", use_container_width=True):
                    try:
                        run_async(call_tool_async("atualizar_sessao_meta", session_id=sid, name=new_name))
                        if sid == st.session_state.session_id:
                            st.session_state.session_name = new_name
                        st.session_state.sessions_list = None
                        st.session_state[f"menu_open_{sid}"] = False
                        st.rerun()
                    except Exception as e:
                        st.error(f"Erro: {e}")

                pin_label = "📌 Desafixar" if pinned else "📌 Fixar no topo"
                if st.button(pin_label, key=f"pin_{sid}", use_container_width=True):
                    try:
                        run_async(call_tool_async(
                            "atualizar_sessao_meta",
                            session_id=sid,
                            pinned="false" if pinned else "true",
                        ))
                        st.session_state.sessions_list = None
                        st.session_state[f"menu_open_{sid}"] = False
                        st.rerun()
                    except Exception as e:
                        st.error(f"Erro: {e}")

                # Exportar
                try:
                    result_load = run_async(call_tool_async("carregar_sessao", session_id=sid))
                    raw_load    = getattr(result_load, "data", getattr(result_load, "output", str(result_load)))
                    data_load   = json.loads(str(raw_load))
                    msgs_export = data_load.get("messages", []) if "error" not in data_load else []
                except Exception:
                    msgs_export = []

                if msgs_export:
                    txt_bytes = _export_txt(sname, msgs_export)
                    st.download_button(
                        "⬇️ Exportar TXT",
                        data=txt_bytes,
                        file_name=f"{sname.replace(' ', '_')}.txt",
                        mime="text/plain",
                        key=f"exp_txt_{sid}",
                        use_container_width=True,
                    )
                    # Garante python-docx disponível
                    try:
                        from docx import Document  # noqa
                    except ImportError:
                        import subprocess, sys
                        subprocess.run([sys.executable, "-m", "pip", "install",
                                        "python-docx", "--quiet"], check=False)
                    docx_bytes = _export_docx(sname, msgs_export)
                    st.download_button(
                        "⬇️ Exportar DOCX",
                        data=docx_bytes,
                        file_name=f"{sname.replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"exp_docx_{sid}",
                        use_container_width=True,
                    )

                st.markdown("<hr style='border-color:rgba(255,255,255,0.25);margin:8px 0;'>", unsafe_allow_html=True)
                if st.button("🗑️ Excluir sessão", key=f"del_sess_{sid}", use_container_width=True):
                    try:
                        run_async(call_tool_async("deletar_sessao", session_id=sid))
                        if sid == st.session_state.session_id:
                            st.session_state.messages            = [{"role": "assistant", "content": "Olá! Como posso ajudar você hoje?"}]
                            st.session_state.session_id          = None
                            st.session_state.session_name        = ""
                            st.session_state.show_confirm_button = False
                        st.session_state.sessions_list = None
                        st.session_state[f"menu_open_{sid}"] = False
                        st.rerun()
                    except Exception as e:
                        st.error(f"Erro: {e}")

                st.markdown("</div>", unsafe_allow_html=True)

st.sidebar.markdown("---")

# ── Nome da sessão ativa (exibido acima das mensagens) ──────────────────────
if st.session_state.session_name:
    col_title, col_edit = st.columns([8, 2])
    with col_title:
        st.markdown(
            f"<h4 style='color:#2E3A8C; margin-bottom:0;'>💬 {st.session_state.session_name}</h4>",
            unsafe_allow_html=True,
        )
    with col_edit:
        if st.button("✏️ Renomear", key="rename_active_top"):
            st.session_state.session_name_editing = not st.session_state.get("session_name_editing", False)

    if st.session_state.get("session_name_editing"):
        new_active_name = st.text_input(
            "Novo nome:", value=st.session_state.session_name,
            key="active_name_input", max_chars=60,
        )
        col_s, col_c = st.columns(2)
        with col_s:
            if st.button("💾 Salvar", key="save_active_name"):
                if new_active_name and new_active_name != st.session_state.session_name:
                    try:
                        if st.session_state.session_id:
                            run_async(call_tool_async(
                                "atualizar_sessao_meta",
                                session_id=st.session_state.session_id,
                                name=new_active_name,
                            ))
                        st.session_state.session_name    = new_active_name
                        st.session_state.sessions_list   = None
                    except Exception as e:
                        st.error(f"Erro ao renomear: {e}")
                st.session_state.session_name_editing = False
                st.rerun()
        with col_c:
            if st.button("❌ Cancelar", key="cancel_active_name"):
                st.session_state.session_name_editing = False
                st.rerun()
    st.markdown("---")

# Renderiza histórico
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

# Entrada principal do chat
if prompt := st.chat_input("Qual a sua dúvida sobre os documentos?"):
    # Ao iniciar uma nova pergunta, escondemos o expander de feedback (ele só volta após avaliação)
    st.session_state.show_feedback_expander = False
    st.session_state.show_confirm_button    = False
    st.session_state.feedback_saved         = False
    st.session_state.feedback_mode          = "partial"
    st.session_state.rerank_scores          = {}

    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user", avatar="👤"):
        st.markdown(prompt)

    with st.spinner("Buscando informações e pensando... 🧠 para gerar a resposta..."):
        try:
            # [MELHORIA] Envia histórico de conversa para contexto multi-turno.
            # Exclui a mensagem atual (ainda não respondida) e filtra só user/assistant.
            history_to_send = [
                {"role": m["role"], "content": m["content"]}
                for m in st.session_state.messages[:-1]  # exclui a msg recém-adicionada
                if m["role"] in ("user", "assistant")
            ][-6:]  # últimas 6 mensagens (3 trocas)

            response = run_async(call_tool_async(
                'ask_question',
                query=prompt,
                force_fresh=bool(st.session_state.get("force_fresh", False)),
                chat_history=history_to_send
            ))
            resp_msg = getattr(response, "data", getattr(response, "output", str(response)))
        except ConnectionRefusedError:
            st.error("⚠️ Servidor MCP fora do ar. Inicie o servidor com: `python mcp_server.py`")
            st.stop()
        except RuntimeError as e:
            err_str = str(e).lower()
            if "connection closed" in err_str or "falhou após" in err_str or "timeout" in err_str:
                st.error(
                    "⏱️ O servidor demorou muito para responder ou a conexão foi encerrada.\n\n"
                    "Isso pode ocorrer quando há muitos documentos e o reranking leva mais tempo "
                    f"que o esperado.\n\n**Detalhe técnico:** {e}"
                )
                st.info("💡 Tente reformular a pergunta de forma mais específica, ou aguarde e tente novamente.")
            else:
                st.error(f"❌ Erro ao se comunicar com o servidor: {e}")
            # Remove a pergunta do histórico (não foi respondida)
            if st.session_state.messages and st.session_state.messages[-1]["role"] == "user":
                st.session_state.messages.pop()
            st.stop()
        except Exception as e:
            st.error(f"❌ Erro inesperado: {type(e).__name__}: {e}")
            if st.session_state.messages and st.session_state.messages[-1]["role"] == "user":
                st.session_state.messages.pop()
            st.stop()

        # Garante que resp_msg é sempre string (evita crash no .strip()/.lower() abaixo)
        if not isinstance(resp_msg, str):
            resp_msg = str(resp_msg)

        # [MELHORIA] Extrai e remove o marcador __RERANK_SCORES__ da resposta
        # (é um dado interno para o feedback, não deve aparecer ao usuário)
        import json as _json
        rerank_scores = {}
        rerank_marker = "__RERANK_SCORES__:"
        if rerank_marker in resp_msg:
            try:
                main_part, scores_part = resp_msg.rsplit(rerank_marker, 1)
                rerank_scores = _json.loads(scores_part.strip())
                resp_msg = main_part.strip()
            except Exception:
                resp_msg = resp_msg.replace(rerank_marker, "").strip()
        st.session_state.rerank_scores = rerank_scores

        with st.chat_message("assistant", avatar="🤖"):            
            # Verifica se é resposta do cache (comparação robusta case-insensitive)
            is_cache_response = resp_msg.strip().lower().startswith("(resposta do cache)")
            if is_cache_response:
                # Remove o prefixo de cache mantendo a capitalização original
                clean_response = re.sub(r"^\(resposta do cache\)\n?", "", resp_msg, flags=re.IGNORECASE)

                if "📚 **Fontes consultadas:**" in clean_response:
                    main_response, sources_section = clean_response.split("📚 **Fontes consultadas:**", 1)
                    st.markdown(main_response)
                    with st.container():
                        st.markdown("📚 **Fontes consultadas:**")
                        st.info(sources_section.strip())
                else:
                    st.markdown(clean_response)

                st.success("💾 Resposta recuperada do cache de perguntas anteriores")
            else:
                if "📚 **Fontes consultadas:**" in resp_msg:
                    main_response, sources_section = resp_msg.split("📚 **Fontes consultadas:**", 1)
                    st.markdown(main_response)
                    with st.container():
                        st.markdown("📚 **Fontes consultadas:**")
                        # [MELHORIA] Exibe legenda dos indicadores de confiança
                        st.caption("⭐⭐⭐ Alta relevância  |  ⭐⭐ Relevância moderada  |  ⭐ Baixa relevância")
                        st.info(sources_section.strip())
                else:
                    st.markdown(resp_msg)

        st.session_state.messages.append({"role": "assistant", "content": resp_msg})

        # Reset da flag force_fresh após uso
        if st.session_state.get("force_fresh", False):
            st.session_state.force_fresh = False

        # ── Auto-save da sessão ──────────────────────────────────────────────
        # Na primeira resposta da sessão, gera nome via LLM em background.
        # Nas subsequentes, apenas atualiza o histórico sem chamar o LLM.
        user_msgs = [m for m in st.session_state.messages if m["role"] == "user"]
        if len(user_msgs) == 1 and not st.session_state.session_name:
            # Primeira pergunta — gera nome sugerido pelo LLM
            try:
                name_result = run_async(call_tool_async(
                    "gerar_nome_sessao",
                    primeira_pergunta=prompt,
                ))
                suggested = getattr(name_result, "data", getattr(name_result, "output", str(name_result)))
                suggested = str(suggested).strip()
                if not suggested or suggested.lower() in ("nova sessão", "nova sessao"):
                    suggested = prompt[:50]
                st.session_state.session_name = suggested
            except Exception:
                st.session_state.session_name = prompt[:50]

        _save_current_session()
        # ── Fim auto-save ────────────────────────────────────────────────────

        # Só exibe o botão se NÃO for uma resposta do cache
        if not is_cache_response:
            st.session_state.ultima_pergunta = prompt
            st.session_state.ultima_resposta = resp_msg
            st.session_state.ultima_fontes = extract_sources_from_response(resp_msg)

            # [FIX-2] Usa prefixo único "não encontrei" que cobre TODAS as variantes
            # de resposta vazia do servidor (gate pós-reranking, fallback final,
            # filtro por tipo). Antes só "não encontrei informação relevante" era
            # verificado, deixando "não encontrei informações suficientes" passar
            # e permitindo que o usuário cacheasse uma resposta de falha como ✅ Correta.
            _resp_lower = resp_msg.strip().lower()
            _is_not_found = (
                _resp_lower.startswith("não encontrei")
                or _resp_lower.startswith("nao encontrei")
                or "❌" in resp_msg[:10]  # erros de sistema (ex.: falha no Ollama)
            )
            if not _is_not_found:
                st.session_state.show_confirm_button = True
            else:
                st.session_state.show_confirm_button = False
        else:
            st.session_state.show_confirm_button = False

# -----------------------------------------------------------------------------
# [MELHORIA] Feedback humano simplificado — 3 botões + modo avançado oculto
# -----------------------------------------------------------------------------
# A interface anterior expunha "prefer_document_types", "must_keywords",
# "query_rewrite" etc. diretamente ao usuário — vocabulário de RAG Engineering
# que confunde usuários finais e reduz a taxa de uso do feedback.
#
# Nova abordagem:
#   Camada 1 (sempre visível): 3 botões de avaliação simples.
#   Camada 2 (após ⚠️/❌):    campo de texto livre opcional.
#   Camada 3 (oculta):         expander "⚙️ Avançado" com o formulário completo.
#
# O mapeamento para prefer/avoid é inferido automaticamente dos rerank_scores.
# -----------------------------------------------------------------------------

if (
    st.session_state.get("ultima_pergunta")
    and st.session_state.get("ultima_resposta")
    and st.session_state.get("show_confirm_button", False)
):
    # [FIX-5] Dupla verificação: show_confirm_button já é False para respostas
    # "não encontradas" (Fix 2), mas esta guarda adicional protege contra
    # qualquer caminho de código que possa ter setado o flag incorretamente.
    _ultima = st.session_state.ultima_resposta.strip().lower()
    _bloqueado = _ultima.startswith("não encontrei") or _ultima.startswith("nao encontrei") or "❌" in st.session_state.ultima_resposta[:10]
    if _bloqueado:
        st.session_state.show_confirm_button = False  # corrige estado inconsistente

if (
    st.session_state.get("ultima_pergunta")
    and st.session_state.get("ultima_resposta")
    and st.session_state.get("show_confirm_button", False)
):
    st.divider()
    st.markdown("**Como foi essa resposta?**")
    fb_col1, fb_col2, fb_col3 = st.columns(3)

    with fb_col1:
        if st.button("✅ Correta", key="fb_correct", use_container_width=True):
            # Registra silenciosamente: reforça as fontes com score alto
            prefer_srcs = [
                src for src, score in st.session_state.get("rerank_scores", {}).items()
                if score >= 7.0
            ]
            prefer_types_inferred = list(set(extract_types(
                st.session_state.get("ultima_fontes") or
                extract_sources_from_response(st.session_state.ultima_resposta)
            )))
            try:
                run_async(call_tool_async(
                    "confirmar_resposta",
                    pergunta=st.session_state.ultima_pergunta,
                    resposta=st.session_state.ultima_resposta
                ))
                if prefer_srcs:
                    run_async(call_tool_async(
                        "registrar_feedback",
                        pergunta=st.session_state.ultima_pergunta,
                        prefer_sources=prefer_srcs,
                        prefer_document_types=prefer_types_inferred,
                        note="auto: resposta marcada como correta"
                    ))
                st.success("✅ Ótimo! Resposta salva como referência.")
                st.session_state.show_confirm_button = False
                st.session_state.show_feedback_expander = False
            except Exception as e:
                st.error(f"Erro ao salvar: {e}")

    with fb_col2:
        if st.button("⚠️ Incompleta", key="fb_partial", use_container_width=True):
            st.session_state.show_confirm_button = False
            st.session_state.show_feedback_expander = True
            st.session_state.feedback_mode = "partial"
            st.rerun()

    with fb_col3:
        if st.button("❌ Errada", key="fb_wrong", use_container_width=True):
            st.session_state.show_confirm_button = False
            st.session_state.show_feedback_expander = True
            st.session_state.feedback_mode = "wrong"
            st.session_state.force_fresh = True
            st.rerun()

# -----------------------------------------------------------------------------
# Painel de feedback expandido (aparece após ⚠️ ou ❌)
# -----------------------------------------------------------------------------
if (
    st.session_state.get("ultima_pergunta")
    and st.session_state.get("ultima_resposta")
    and st.session_state.get("show_feedback_expander", False)
):
    feedback_mode = st.session_state.get("feedback_mode", "partial")
    mode_label    = "incompleta" if feedback_mode == "partial" else "errada"
    mode_icon     = "⚠️" if feedback_mode == "partial" else "❌"

    st.markdown(f"**{mode_icon} Você marcou a resposta como {mode_label}.**")

    free_text = st.text_area(
        "O que estava faltando ou errado? (opcional — ajuda o sistema a melhorar)",
        placeholder="Ex.: 'Não mencionou a etapa de aprovação do gerente' ou 'Usou um relatório antigo, deveria usar o manual atual'",
        height=80,
        key="fb_free_text"
    )

    fb_submit_col, fb_retry_col = st.columns(2)

    with fb_submit_col:
        if st.button("💾 Enviar feedback", key="fb_submit_simple", use_container_width=True):
            with st.spinner("Processando feedback..."):
                try:
                    rerank_scores = st.session_state.get("rerank_scores", {})

                    # Usa a função centralizada: detecta negações, instruções de
                    # formato e extrai prefer/avoid/must_keywords/response_instruction.
                    prefer_srcs, avoid_srcs, avoid_doc_types, must_kws, resp_instruction, explicit_prefer_srcs = _infer_feedback_from_text(
                        free_text=free_text,
                        rerank_scores=rerank_scores,
                        mode=feedback_mode
                    )

                    # No modo 'wrong', penaliza adicionalmente fontes com score
                    # baixo mesmo que o usuário não as tenha nomeado explicitamente
                    if feedback_mode == "wrong" and not avoid_srcs:
                        avoid_srcs = [s for s, sc in rerank_scores.items() if sc < 5.0]

                    note = (
                        f"{'parcial' if feedback_mode == 'partial' else 'errada'}: {free_text}"
                        if free_text else
                        f"auto: resposta marcada como {'incompleta' if feedback_mode == 'partial' else 'errada'}"
                    )

                    run_async(call_tool_async(
                        "registrar_feedback",
                        pergunta=st.session_state.ultima_pergunta,
                        prefer_sources=prefer_srcs,
                        avoid_sources=avoid_srcs,
                        avoid_document_types=avoid_doc_types,
                        explicit_prefer_sources=explicit_prefer_srcs,
                        must_keywords=must_kws,
                        response_instruction=resp_instruction,
                        note=note
                    ))
                    st.success("Feedback registrado! O sistema vai usar isso em perguntas similares.")
                    st.session_state.feedback_saved = True
                    st.session_state.show_feedback_expander = False
                    st.session_state.knowledge_stats = None  # força atualização do painel
                except Exception as e:
                    st.error(f"❌ Erro ao salvar feedback: {e}")

    with fb_retry_col:
        if st.session_state.get("feedback_saved", False):
            if st.button("🔁 Tentar novamente", key="fb_retry_simple", use_container_width=True):
                st.session_state.force_fresh = True
                st.session_state.rerun_after_feedback = True
                st.rerun()

    # --- Modo avançado oculto por padrão ---
    with st.expander("⚙️ Configuração avançada de retrieval"):
        st.write(
            "Configure manualmente quais documentos/tipos priorizar ou evitar "
            "em perguntas similares no futuro."
        )

        if st.session_state.get("feedback_toast"):
            msg = st.session_state.feedback_toast
            try:
                st.toast(msg)
            except Exception:
                pass
            st.success(msg)
            st.session_state.feedback_toast = None

        sources_lines = (
            st.session_state.get("ultima_fontes")
            or extract_sources_from_response(st.session_state.ultima_resposta)
        )
        filenames_only = extract_filenames(sources_lines)
        types_found    = extract_types(sources_lines)

        if "all_types_cache" not in st.session_state:
            st.session_state.all_types_cache = []
        if "all_sources_cache" not in st.session_state:
            st.session_state.all_sources_cache = []

        if not st.session_state.all_types_cache:
            try:
                resp_types = run_async(call_tool_async("list_document_types"))
                all_types  = getattr(resp_types, "data", getattr(resp_types, "output", []))
                if isinstance(all_types, str):
                    all_types = [all_types]
                if not isinstance(all_types, list):
                    all_types = []
                st.session_state.all_types_cache = all_types
            except Exception:
                pass

        if not st.session_state.all_sources_cache:
            try:
                resp_sources = run_async(call_tool_async("list_sources"))
                all_sources  = getattr(resp_sources, "data", getattr(resp_sources, "output", []))
                if isinstance(all_sources, str):
                    all_sources = [all_sources]
                if not isinstance(all_sources, list):
                    all_sources = []
                st.session_state.all_sources_cache = all_sources
            except Exception:
                pass

        all_types    = st.session_state.all_types_cache
        all_sources  = st.session_state.all_sources_cache
        type_options = all_types   if all_types   else (types_found   if types_found   else doc_type_options)
        source_options = all_sources if all_sources else filenames_only

        with st.form("feedback_form_advanced"):
            prefer_types = st.multiselect("Preferir tipos de documento", options=type_options,
                                          default=st.session_state.get("fb_prefer_types", []), key="fb_prefer_types")
            avoid_types  = st.multiselect("Evitar tipos de documento",   options=type_options,
                                          default=st.session_state.get("fb_avoid_types",  []), key="fb_avoid_types")
            prefer_sources = st.multiselect("Preferir estes arquivos",   options=source_options,
                                            default=st.session_state.get("fb_prefer_sources", []), key="fb_prefer_sources")
            avoid_sources  = st.multiselect("Evitar estes arquivos",     options=source_options,
                                            default=st.session_state.get("fb_avoid_sources",  []), key="fb_avoid_sources")
            must_keywords_text = st.text_input("Must-keywords (separadas por vírgula)",
                                               value=st.session_state.get("fb_must_keywords_text", ""),
                                               key="fb_must_keywords_text",
                                               help="Termos/siglas obrigatórias para ampliar a busca.")
            response_instruction_adv = st.text_area(
                "Instrução de formato (como o modelo deve responder)",
                value=st.session_state.get("fb_response_instruction", ""),
                key="fb_response_instruction",
                height=70,
                help="Ex.: 'Responda em uma frase objetiva.' ou 'Use tópicos numerados.' ou 'Seja conciso.'",
            )
            query_rewrite = st.text_area("Query rewrite (opcional)",
                                         value=st.session_state.get("fb_query_rewrite", ""),
                                         key="fb_query_rewrite", height=70)
            note = st.text_area("Nota (opcional)", value=st.session_state.get("fb_note", ""),
                                key="fb_note", height=80)
            submitted_adv = st.form_submit_button("💾 Salvar configuração avançada")

        if submitted_adv:
            with st.spinner("Salvando..."):
                try:
                    must_keywords = [x.strip() for x in (must_keywords_text or "").split(",") if x.strip()]

                    # [FIX-ADV-2] prefer_sources selecionado manualmente no formulário
                    # avançado é tratado como HARD (explicit_prefer_sources), não soft.
                    # O usuário que escolhe explicitamente um arquivo no multiselect
                    # espera que apenas aquele arquivo seja usado — comportamento HARD.
                    explicit_prefer_adv = list(prefer_sources) if prefer_sources else []

                    # [FIX-ADV-1] Campo 'rule' removido da UI — não tinha parâmetro
                    # correspondente no servidor e era ignorado silenciosamente.
                    # O campo 'note' cobre o mesmo caso de uso de forma funcional.

                    fb_result = run_async(call_tool_async(
                        "registrar_feedback",
                        pergunta=st.session_state.ultima_pergunta,
                        prefer_document_types=prefer_types,
                        avoid_document_types=avoid_types,
                        prefer_sources=prefer_sources,
                        avoid_sources=avoid_sources,
                        explicit_prefer_sources=explicit_prefer_adv,
                        must_keywords=must_keywords,
                        response_instruction=(response_instruction_adv or "").strip(),
                        query_rewrite=(query_rewrite or "").strip(),
                        note=note
                    ))
                    fb_msg = getattr(fb_result, "data", getattr(fb_result, "output", str(fb_result)))
                    st.session_state.feedback_toast = fb_msg
                    st.success(fb_msg)
                    st.session_state.feedback_saved = True
                    st.session_state.force_fresh = True
                    # [FIX-ADV-3] Aciona retry automático igual ao formulário simples
                    st.session_state.rerun_after_feedback = True
                    st.session_state.knowledge_stats = None
                except Exception as e:
                    st.error(f"❌ Erro ao salvar: {e}")



# -----------------------------------------------------------------------------
# NOVO: retentar a MESMA pergunta após o usuário ensinar (feedback salvo)
# -----------------------------------------------------------------------------
if st.session_state.get("rerun_after_feedback", False) and st.session_state.get("ultima_pergunta"):
    retry_query = st.session_state.ultima_pergunta

    with st.spinner("🔁 Tentando novamente com o novo ensino (ignorando cache)..."):
        try:
            # [FIX-BUG3] Aguarda o Qdrant processar o upsert do feedback antes
            # de disparar a retry. O wait=True no upsert garante que a operação
            # foi confirmada pelo servidor Qdrant, mas há latência de indexação
            # do vetor em memória. 0.5s é suficiente para o índice HNSW ser atualizado.
            import time as _time
            _time.sleep(0.5)

            retry_response = run_async(call_tool_async(
                "ask_question",
                query=retry_query,
                force_fresh=True
            ))
            retry_msg = getattr(retry_response, "data", getattr(retry_response, "output", str(retry_response)))
            if not isinstance(retry_msg, str):
                retry_msg = str(retry_msg)

            # Mostra somente a nova resposta do assistente (sem repetir a pergunta)
            with st.chat_message("assistant", avatar="🤖"):
                if "📚 **Fontes consultadas:**" in retry_msg:
                    main_response, sources_section = retry_msg.split("📚 **Fontes consultadas:**", 1)
                    st.markdown(main_response)
                    with st.container():
                        st.markdown("📚 **Fontes consultadas:**")
                        st.info(sources_section.strip())
                else:
                    st.markdown(retry_msg)

            # Guarda histórico e sobrescreve "última resposta" (para permitir avaliar novamente)
            st.session_state.messages.append({"role": "assistant", "content": retry_msg})
            st.session_state.ultima_resposta = retry_msg

        except ConnectionRefusedError:
            st.error("⚠️ Servidor MCP fora do ar ao retentar. Inicie com: `python mcp_server.py`")
        except RuntimeError as e:
            err_str = str(e).lower()
            if "connection closed" in err_str or "falhou após" in err_str or "timeout" in err_str:
                st.error(
                    "⏱️ A retentativa também sofreu timeout ou conexão fechada.\n\n"
                    "O servidor pode estar sobrecarregado com o reranking. "
                    "Aguarde alguns segundos e tente novamente."
                )
            else:
                st.error(f"❌ Erro ao retentar pergunta: {e}")
        except Exception as e:
            st.error(f"❌ Erro inesperado ao retentar: {type(e).__name__}: {e}")

    # Consome a flag (evita loop)
    st.session_state.rerun_after_feedback = False