# -*- coding: utf-8 -*-
# Arquivo: api/utils/report_generator.py
"""
Módulo para geração de relatórios inteligentes baseados em conversas.
Analisa o histórico de perguntas e respostas para gerar insights e recomendações.
"""

import io
import json
import re
from datetime import datetime
from typing import List, Dict, Any, Optional

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


class ReportGenerator:
    """
    Gera relatórios inteligentes a partir do histórico de conversas.
    """

    # Limites para colocar tabela no relatório (profissional e sem explodir o DOCX)
    TABLE_MAX_COLS = 6
    TABLE_MAX_ROWS = 30
    TABLE_CELL_MAX_CHARS = 200  # evita células gigantes

    def __init__(self, vanna_instance):
        """
        Inicializa o gerador de relatórios.

        Args:
            vanna_instance: Instância do Vanna para geração de insights com LLM
        """
        self.vanna = vanna_instance

    # ---------------------------------------------------------------------
    # Conversa -> análise estruturada
    # ---------------------------------------------------------------------
    def analyze_conversation(self, messages: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        Analisa uma conversa completa e extrai insights.

        Args:
            messages: Lista de mensagens da conversa

        Returns:
            Dicionário com análise da conversa
        """
        analysis = {
            "total_messages": len(messages),
            "user_questions": [],
            "assistant_responses": [],
            "queries_executed": [],
            "data_insights": [],
            "topics": set(),
            "timestamp": datetime.now().isoformat(),
        }

        for msg in messages:
            if msg.get("role") == "user":
                analysis["user_questions"].append(msg.get("content", ""))
                self._extract_topics(msg.get("content", ""), analysis["topics"])

            elif msg.get("role") == "assistant":
                content = msg.get("content")

                # Se for resposta estruturada (dict) vinda do backend
                if isinstance(content, dict):
                    df_preview = self._make_dataframe_preview(content.get("dataframe"))

                    analysis["assistant_responses"].append(
                        {
                            "summary": content.get("summary", ""),
                            "sql": content.get("sql", ""),
                            "has_data": bool(content.get("dataframe")),
                            "has_chart": bool(content.get("chart")),
                            "from_memory": content.get("from_memory", False),
                            "dataframe_preview": df_preview,  # <-- usado para tabela
                        }
                    )

                    if content.get("sql"):
                        analysis["queries_executed"].append(
                            {
                                "sql": content.get("sql", ""),
                                "summary": content.get("summary", ""),
                                "row_count": len(content.get("dataframe", [])) if content.get("dataframe") else 0,
                                "dataframe_preview": df_preview,
                            }
                        )

                    if content.get("dataframe"):
                        self._extract_data_insights(content, analysis["data_insights"])

        analysis["topics"] = list(analysis["topics"])
        return analysis

    def _extract_topics(self, question: str, topics: set):
        """Extrai tópicos principais de uma pergunta."""
        q_lower = (question or "").lower()

        topic_keywords = {
            "clientes": ["cliente", "customer", "company"],
            "vendas": ["venda", "vender", "receita", "faturamento"],
            "produtos": ["produto", "item", "mercadoria"],
            "pedidos": ["pedido", "order", "solicitação", "request"],
            "status": ["status", "estado", "situação"],
            "tempo": ["mês", "ano", "data", "período", "dia", "semana"],
            "localização": ["cidade", "estado", "região", "local"],
            "quantidade": ["total", "quantidade", "count", "número", "quantos"],
            "análise": ["top", "ranking", "maior", "menor", "melhor", "pior"],
            "garantia": ["garantia", "warranty", "rma"],
            "devolução": ["devolução", "return"],
        }

        for topic, keywords in topic_keywords.items():
            if any(kw in q_lower for kw in keywords):
                topics.add(topic)

    def _extract_data_insights(self, response: Dict[str, Any], insights: List[str]):
        """Extrai insights dos dados retornados."""
        dataframe = response.get("dataframe", [])
        if not dataframe:
            return

        df = pd.DataFrame(dataframe)

        row_count = len(df)
        if row_count > 0:
            insights.append(f"Análise retornou {row_count} registro(s)")

        numeric_cols = df.select_dtypes(include=["int64", "float64"]).columns
        if len(numeric_cols) > 0:
            for col in numeric_cols:
                total = df[col].sum()
                if total > 0:
                    insights.append(f"{col}: total de {total:,.0f}")

    # ---------------------------------------------------------------------
    # LLM insights
    # ---------------------------------------------------------------------
    def generate_llm_insights(self, analysis: Dict[str, Any]) -> str:
        """Usa o LLM para gerar insights e recomendações baseadas na análise."""
        context = self._build_llm_context(analysis)

        prompt = f"""Analise a seguinte conversa e gere um resumo executivo com insights e recomendações:

CONTEXTO DA CONVERSA:
{context}

Por favor, forneça:

1. **RESUMO EXECUTIVO** (2-3 parágrafos):
   - Principais tópicos abordados
   - Padrões identificados nos dados

2. **INSIGHTS PRINCIPAIS** (3-5 pontos):
   - Descobertas importantes
   - Tendências ou anomalias

3. **RECOMENDAÇÕES** (3-5 ações):
   - Ações concretas baseadas nos dados
   - Próximos passos sugeridos

Seja objetivo, direto e focado em apoiar tomada de decisão.
"""

        try:
            # usa Ollama direto (como no seu original)
            from ollama import Client

            client = Client()
            response = client.generate(
                model=self.vanna.config.get("model", "gpt-oss:20b"),
                prompt=prompt,
                options={"num_predict": 1024},
            )
            return response.get("response", "")
        except Exception as e:
            print(f"⚠️ Erro ao gerar insights com LLM: {e}")
            return self._generate_fallback_insights(analysis)

    def _build_llm_context(self, analysis: Dict[str, Any]) -> str:
        """Constrói contexto textual para o LLM."""
        context_parts = []

        if analysis["user_questions"]:
            context_parts.append("PERGUNTAS DO USUÁRIO:")
            for i, q in enumerate(analysis["user_questions"][:10], 1):
                context_parts.append(f"{i}. {q}")

        if analysis["topics"]:
            context_parts.append(f"\nTÓPICOS ABORDADOS: {', '.join(analysis['topics'])}")

        if analysis["queries_executed"]:
            context_parts.append(f"\nCONSULTAS EXECUTADAS: {len(analysis['queries_executed'])}")
            for query in analysis["queries_executed"][:5]:
                context_parts.append(f"- {query['summary']} ({query['row_count']} registros)")

        if analysis["data_insights"]:
            context_parts.append("\nDADOS IDENTIFICADOS:")
            for insight in analysis["data_insights"][:10]:
                context_parts.append(f"- {insight}")

        return "\n".join(context_parts)

    def _generate_fallback_insights(self, analysis: Dict[str, Any]) -> str:
        """Gera insights básicos sem LLM (fallback)."""
        insights = []

        insights.append("# RESUMO EXECUTIVO\n")
        insights.append(
            f"Esta conversa abordou {len(analysis['user_questions'])} perguntas sobre os seguintes tópicos: "
            f"{', '.join(analysis['topics'])}.\n"
        )
        insights.append(f"Foram executadas {len(analysis['queries_executed'])} consultas ao banco de dados.\n")

        insights.append("\n# INSIGHTS PRINCIPAIS\n")
        for i, insight in enumerate(analysis["data_insights"][:5], 1):
            insights.append(f"{i}. {insight}")

        insights.append("\n\n# RECOMENDAÇÕES\n")
        insights.append("1. Revisar os dados identificados para validar padrões")
        insights.append("2. Aprofundar análise nos tópicos mais consultados")
        insights.append("3. Monitorar métricas identificadas periodicamente")

        return "\n".join(insights)

    # ---------------------------------------------------------------------
    # Markdown simples -> DOCX
    # ---------------------------------------------------------------------
    def _strip_md(self, text: str) -> str:
        if not text:
            return ""
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"__(.+?)__", r"\1", text)
        text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", text)
        text = re.sub(r"(?<!_)_(?!_)(.+?)(?<!_)_(?!_)", r"\1", text)
        return text

    def _add_inline_bold_runs(self, paragraph, line: str):
        """Preserva trechos **...** como bold."""
        if not line:
            return
        parts = re.split(r"(\*\*.+?\*\*)", line)
        for part in parts:
            if not part:
                continue
            m = re.match(r"^\*\*(.+?)\*\*$", part)
            if m:
                run = paragraph.add_run(m.group(1))
                run.bold = True
            else:
                paragraph.add_run(self._strip_md(part))

    def _add_markdown_block(self, doc: Document, md_text: str):
        """
        Converte markdown simples em conteúdo docx:
        - Headings com #
        - Headings em **TÍTULO**
        - Listas: 1. ... | - ... | • ...
        - Separador: --- / *** / ___
        - Negrito inline: **texto**
        """
        if not md_text:
            return

        for raw in md_text.splitlines():
            line = raw.rstrip("\n")
            stripped = line.strip()

            if not stripped:
                doc.add_paragraph("")
                continue

            if stripped in ("---", "***", "___"):
                p = doc.add_paragraph()
                run = p.add_run("―" * 40)
                run.font.size = Pt(10)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

            if stripped.startswith("#"):
                level = min(stripped.count("#"), 4)
                title = stripped.lstrip("#").strip()
                doc.add_heading(self._strip_md(title), level)
                continue

            m_heading_bold = re.match(r"^\*\*(.+?)\*\*$", stripped)
            if m_heading_bold:
                doc.add_heading(self._strip_md(m_heading_bold.group(1)), 2)
                continue

            m_num = re.match(r"^\d+\.\s+(.*)$", stripped)
            if m_num:
                p = doc.add_paragraph(style="List Number")
                self._add_inline_bold_runs(p, m_num.group(1))
                continue

            m_bul = re.match(r"^[-•]\s+(.*)$", stripped)
            if m_bul:
                p = doc.add_paragraph(style="List Bullet")
                self._add_inline_bold_runs(p, m_bul.group(1))
                continue

            p = doc.add_paragraph()
            self._add_inline_bold_runs(p, stripped)

    # ---------------------------------------------------------------------
    # DataFrame -> tabela DOCX (quando pequeno)
    # ---------------------------------------------------------------------
    def _make_dataframe_preview(self, dataframe: Any) -> Optional[List[Dict[str, Any]]]:
        """
        Retorna preview do dataframe para tabelas no DOCX, ou None se não for adequado.
        Espera dataframe como list[dict] (que é o que sua API retorna).
        """
        if not dataframe or not isinstance(dataframe, list):
            return None
        if len(dataframe) == 0:
            return None
        if not isinstance(dataframe[0], dict):
            return None

        try:
            df = pd.DataFrame(dataframe)
            if df.empty:
                return None

            # Limites para tabela "bonita"
            if df.shape[1] > self.TABLE_MAX_COLS:
                return None

            df_small = df.head(self.TABLE_MAX_ROWS).copy()

            # Converte NaN/NaT e trunca strings longas
            df_small = df_small.where(pd.notnull(df_small), "")

            for c in df_small.columns:
                df_small[c] = df_small[c].astype(str).map(lambda s: s[: self.TABLE_CELL_MAX_CHARS])

            return df_small.to_dict(orient="records")
        except Exception:
            return None

    def _add_dataframe_table(self, doc: Document, rows: List[Dict[str, Any]], title: Optional[str] = None):
        """Adiciona uma tabela no DOCX a partir de list[dict]."""
        if not rows:
            return
        columns = list(rows[0].keys())
        if not columns:
            return

        if title:
            doc.add_paragraph(title).runs[0].bold = True

        table = doc.add_table(rows=len(rows) + 1, cols=len(columns))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Cabeçalho
        hdr_cells = table.rows[0].cells
        for j, col in enumerate(columns):
            hdr_cells[j].text = str(col)
            for run in hdr_cells[j].paragraphs[0].runs:
                run.bold = True

        # Linhas
        for i, r in enumerate(rows, start=1):
            row_cells = table.rows[i].cells
            for j, col in enumerate(columns):
                val = r.get(col, "")
                row_cells[j].text = "" if val is None else str(val)

        # Espaço após tabela
        doc.add_paragraph("")

    # ---------------------------------------------------------------------
    # DOCX Report
    # ---------------------------------------------------------------------
    def generate_docx_report(self, messages: List[Dict[str, Any]], title: str = "Relatório de Análise de Dados") -> io.BytesIO:
        """
        Gera relatório em formato DOCX.

        Args:
            messages: Histórico de mensagens
            title: Título do relatório

        Returns:
            BytesIO com conteúdo do documento
        """
        analysis = self.analyze_conversation(messages)
        llm_insights = self.generate_llm_insights(analysis)

        doc = Document()

        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Data de Geração: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph(f"Total de Mensagens: {analysis['total_messages']}")
        doc.add_paragraph(f"Perguntas Analisadas: {len(analysis['user_questions'])}")
        doc.add_paragraph("")

        # Insights do LLM (com parsing de markdown simples)
        doc.add_heading("Análise Inteligente", 1)
        self._add_markdown_block(doc, llm_insights)

        doc.add_page_break()

        # Tópicos
        doc.add_heading("Tópicos Abordados", 1)
        if analysis["topics"]:
            for topic in analysis["topics"]:
                doc.add_paragraph(topic, style="List Bullet")
        else:
            doc.add_paragraph("Nenhum tópico específico identificado.")

        # Histórico Q/A (+ tabelas quando apropriado)
        doc.add_heading("Histórico de Perguntas", 1)
        for i, question in enumerate(analysis["user_questions"], 1):
            p = doc.add_paragraph()
            p.add_run(f"Q{i}: ").bold = True
            p.add_run(question)

            if i <= len(analysis["assistant_responses"]):
                resp = analysis["assistant_responses"][i - 1]
                doc.add_paragraph(f"Resposta: {resp['summary']}", style="List Bullet 2")

                if resp["has_data"]:
                    doc.add_paragraph("✓ Retornou dados", style="List Bullet 3")

                    # Tabela (se o preview foi gerado)
                    if resp.get("dataframe_preview"):
                        self._add_dataframe_table(
                            doc,
                            resp["dataframe_preview"],
                            title="Amostra de dados (preview):",
                        )

                if resp["has_chart"]:
                    doc.add_paragraph("✓ Gerou gráfico", style="List Bullet 3")

        # Consultas SQL (+ tabelas quando apropriado)
        if analysis["queries_executed"]:
            doc.add_page_break()
            doc.add_heading("Consultas SQL Executadas", 1)

            for i, query in enumerate(analysis["queries_executed"], 1):
                doc.add_heading(f"Consulta {i}", 2)
                doc.add_paragraph(f"Descrição: {query['summary']}")
                doc.add_paragraph(f"Registros retornados: {query['row_count']}")

                # SQL em bloco
                p = doc.add_paragraph()
                run = p.add_run(query["sql"])
                run.font.name = "Courier New"
                run.font.size = Pt(9)

                # Tabela (se couber)
                if query.get("dataframe_preview"):
                    self._add_dataframe_table(doc, query["dataframe_preview"], title="Amostra de resultados (preview):")

                doc.add_paragraph("")

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

    # ---------------------------------------------------------------------
    # JSON Report
    # ---------------------------------------------------------------------
    def generate_json_report(self, messages: List[Dict[str, Any]]) -> str:
        """
        Gera relatório em formato JSON.
        """
        analysis = self.analyze_conversation(messages)
        llm_insights = self.generate_llm_insights(analysis)

        report = {
            "metadata": {
                "title": "Relatório de Análise de Dados",
                "generated_at": datetime.now().isoformat(),
                "total_messages": analysis["total_messages"],
                "total_questions": len(analysis["user_questions"]),
                "total_queries": len(analysis["queries_executed"]),
            },
            "analysis": analysis,
            "llm_insights": llm_insights,
        }

        return json.dumps(report, indent=2, ensure_ascii=False)

